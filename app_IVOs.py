import subprocess
import sys

try:
    import html5lib
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "html5lib"])
    import html5lib

from sentence_transformers import SentenceTransformer, util
import streamlit as st
from docx import Document
from io import BytesIO
import random
import pandas as pd
from datetime import datetime

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}

def formatear_fecha_es(fecha):
    """
    Formatea fechas en español sin depender de locale ni babel.
    Acepta datetime, string o NaT.
    """
    if not fecha or pd.isna(fecha):
        return "Fecha desconocida"

    try:
        if not isinstance(fecha, (datetime, pd.Timestamp)):
            fecha = pd.to_datetime(fecha, errors="coerce")
        if pd.isna(fecha):
            return "Fecha desconocida"

        dia = fecha.day
        mes = MESES_ES.get(fecha.month, "")
        año = fecha.year
        return f"{dia} de {mes} de {año}"
    except Exception:
        return str(fecha)

# CARGA DEL LISTADO DE MIEMBROS
#-------------------------------------------------------------------------------------------
import requests
import pandas as pd

url = "https://secpho.org/wp-json/reports/v1/members?auth=c9J2mL7vT8sW4xAfB3eR6zNpQ1HdUgVKtrXa"

headers = {
    "Accept": "application/json",
    "User-Agent": "Mozilla/5.0"
}

response = requests.get(url, headers=headers)

if response.status_code == 200:
    data = response.json()
    miembros = pd.DataFrame(data)
else:
    miembros = pd.DataFrame()  # Devuelve un DataFrame vacío si hay error
miembros = miembros.T
import unicodedata

# Función para normalizar texto
def normalizar_texto(texto):
    if pd.isna(texto):
        return ''
    texto = str(texto).lower().replace(" ", "")
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("utf-8")
    return texto

# Aplicar la función para crear la columna "Nombre completo"
miembros["Nombre completo"] = (
    miembros["Nombre"].apply(normalizar_texto) +
    miembros["Apellidos"].apply(normalizar_texto)
)

def reemplazar_guion_emisor(valor):
    if isinstance(valor, str):
        return valor.replace(' &#8211;', '-')
    elif isinstance(valor, list):
        return [str(item).replace(' &#8211;', '-') for item in valor]
    else:
        return valor  # Para NaN u otros tipos
def reemplazar_guion_emisorr(valor):
    if isinstance(valor, str):
        return valor.replace('&#038;', '&')
    elif isinstance(valor, list):
        return [str(item).replace('&#038;', '&') for item in valor]
    else:
        return valor  # Para NaN u otros tipos

miembros['Socio'] = miembros['Socio'].apply(reemplazar_guion_emisor)
miembros['Socio'] = miembros['Socio'].apply(reemplazar_guion_emisorr)
#-----------------------------------------------------------------------------------------

# CARGA LISTADO SUSCRIPTORES
#-----------------------------------------------------------------------------------------------------
url = "https://secpho.org/wp-json/reports/v1/suscriptores?auth=c9J2mL7vT8sW4xAfB3eR6zNpQ1HdUgVKtrXa"

headers = {
    "Accept": "application/json",
    "User-Agent": "Mozilla/5.0"
}

response = requests.get(url, headers=headers)

if response.status_code == 200:
    data = response.json()
    suscriptores = pd.DataFrame(data)
else:
    suscriptores = pd.DataFrame()  # Devuelve un DataFrame vacío si hay error
# Función para normalizar texto
def normalizar_texto(texto):
    if pd.isna(texto):
        return ''
    texto = str(texto).lower().replace(" ", "")
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("utf-8")
    return texto

# Aplicar la función para crear la columna "Nombre completo"
suscriptores["Nombre completo"] = (
    suscriptores["Nombre"].apply(normalizar_texto) +
    suscriptores["Apellidos"].apply(normalizar_texto)
)

def reemplazar_guion_emisor(valor):
    if isinstance(valor, str):
        return valor.replace(' &#8211;', '-')
    elif isinstance(valor, list):
        return [str(item).replace(' &#8211;', '-') for item in valor]
    else:
        return valor  # Para NaN u otros tipos
def reemplazar_guion_emisorr(valor):
    if isinstance(valor, str):
        return valor.replace('&#038;', '&')
    elif isinstance(valor, list):
        return [str(item).replace('&#038;', '&') for item in valor]
    else:
        return valor  # Para NaN u otros tipos

suscriptores['Socio'] = suscriptores['Socio'].apply(reemplazar_guion_emisor)
suscriptores['Socio'] = suscriptores['Socio'].apply(reemplazar_guion_emisorr)
#----------------------------------------------------------------------------------------------------------

# CARGA DE LOS SOCIOS
#-----------------------------------------------------------------------------------------------------------
url = "https://secpho.org/wp-json/reports/v1/datosnegocio?auth=c9J2mL7vT8sW4xAfB3eR6zNpQ1HdUgVKtrXa"

headers = {
    "Accept": "application/json",
    "User-Agent": "Mozilla/5.0"
}

response = requests.get(url, headers=headers)

if response.status_code == 200:
    data = response.json()
    socios = pd.DataFrame(data)
else:
    socios = pd.DataFrame()  # Devuelve un DataFrame vacío si hay error

def reemplazar_guion_emisor(valor):
    if isinstance(valor, str):
        return valor.replace(' &#8211;', '-')
    elif isinstance(valor, list):
        return [str(item).replace(' &#8211;', '-') for item in valor]
    else:
        return valor  # Para NaN u otros tipos
def reemplazar_guion_emisorr(valor):
    if isinstance(valor, str):
        return valor.replace('&#038;', '&')
    elif isinstance(valor, list):
        return [str(item).replace('&#038;', '&') for item in valor]
    else:
        return valor  # Para NaN u otros tipos

socios['Socio'] = socios['Socio'].apply(reemplazar_guion_emisor)
socios['Socio'] = socios['Socio'].apply(reemplazar_guion_emisorr)

# Función para normalizar texto
def normalizar_texto(texto):
    if pd.isna(texto):
        return ''
    texto = str(texto).lower().replace(" ", "")
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("utf-8")
    return texto

# Aplicar la función para crear la columna "Nombre completo"
socios["Nombre completo"] = (
    socios["Socio"].apply(normalizar_texto)
)
#-------------------------------------------------------------------------------------------------------------

# CARGA DATOS DE CONTACTO DE SOCIOS
#-------------------------------------------------------------------------------------------------------------
url = "https://secpho.org/wp-json/reports/v1/datoscontacto?auth=c9J2mL7vT8sW4xAfB3eR6zNpQ1HdUgVKtrXa"

headers = {
    "Accept": "application/json",
    "User-Agent": "Mozilla/5.0"
}

response = requests.get(url, headers=headers)

if response.status_code == 200:
    data = response.json()
    contacto_socios = pd.DataFrame(data)
else:
    contacto_socios = pd.DataFrame()  # Devuelve un DataFrame vacío si hay error
# Función para normalizar texto
def normalizar_texto(texto):
    if pd.isna(texto):
        return ''
    texto = str(texto).lower().replace(" ", "")
    texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("utf-8")
    return texto

# Aplicar la función para crear la columna "Nombre completo"
contacto_socios["Nombre completo"] = (
    contacto_socios["Entidad"].apply(normalizar_texto)
)
#------------------------------------------------------------------------------------------------------------

#CARGA DE LISTADO DE EVENTOS PASADOS
#------------------------------------------------------------------------------------------------------------
url = "https://secpho.org/wp-json/reports/v1/actosagenda?auth=c9J2mL7vT8sW4xAfB3eR6zNpQ1HdUgVKtrXa"

headers = {
    "Accept": "application/json",
    "User-Agent": "Mozilla/5.0"
}

response = requests.get(url, headers=headers)

if response.status_code == 200:
    data = response.json()
    eventos_pasados = pd.DataFrame(data)
else:
    eventos_pasados = pd.DataFrame()
eventos_pasados=eventos_pasados.T
# Conjunto de tecnologías que representan "Todas"
tecnologias_completas = {
    'AR/VR', 'Biotecnología', 'Blockchain', 'Ciberseguridad y Cloud Computing',
    'Fabricación Aditiva', 'Fotónica', 'Iluminación avanzada',
    'Inteligencia Artificial y datos', 'IoT', 'Materiales Avanzados',
    'Microelectrónica y Semiconductores', 'Robótica y Drones',
    'Sensórica avanzada', 'Sistemas de Comunicación y Transmisión de Datos',
    'Sistemas láser', 'Tecnologías Cuánticas'
}

# Función que reemplaza la lista por ["Todas"] si contiene exactamente esas 16 tecnologías
def reemplazar_por_todas(tecnologias):
    if isinstance(tecnologias, list) and set(tecnologias) == tecnologias_completas:
        return ["Todas"]
    return tecnologias

# Aplicar la función
eventos_pasados['Tecnología'] = eventos_pasados['Tecnología'].apply(reemplazar_por_todas)
# Conjunto completo de sectores que representan "Todas"
sectores_completos = {
    'Aeronáutica', 'Agroalimentación', 'Automoción', 'Construcción',
    'Cosmética y Estética', 'Defensa', 'Espacio', 'Fabricación industrial',
    'Ferroviario', 'Industria química', 'Logística', 'Materias primas',
    'Medioambiente', 'Naval', 'Patrimonio Cultural', 'Salud',
    'Sector energético', 'Sector farmacéutico', 'Seguridad',
    'Telecomunicaciones', 'Turismo'
}

# Función que reemplaza por ['Todas'] si el contenido es exactamente igual al conjunto
def reemplazar_sector_por_todas(sectores):
    if isinstance(sectores, list) and set(sectores) == sectores_completos:
        return ["Todos"]
    return sectores

# Aplicar la función a la columna 'Sector'
eventos_pasados['Sector'] = eventos_pasados['Sector'].apply(reemplazar_sector_por_todas)
eventos_pasados['Título'] = eventos_pasados['Título'].str.replace(' &#8211;', ':', regex=False)
eventos_pasados['Título'] = eventos_pasados['Título'].str.replace('&#038;', '&', regex=False)
#------------------------------------------------------------------------------------------------------------

# OBTENCIÓN DE LISTADO DE EVENTOS FUTUROS
#------------------------------------------------------------------------------------------------------------
from datetime import datetime

# Asegúrate de que la columna 'fecha' es de tipo datetime
eventos_pasados["Fecha"] = pd.to_datetime(eventos_pasados["Fecha"], errors="coerce")

# Obtener la fecha actual
hoy = pd.Timestamp(datetime.today().date())

# Filtrar eventos cuya fecha sea mayor que hoy
próximos_eventos = eventos_pasados[eventos_pasados["Fecha"] > hoy].copy()
#------------------------------------------------------------------------------------------------------------

#CARGA DE LISTADO DE EVENTOS ASISTIDOS
#------------------------------------------------------------------------------------------------------------
import os

archivos_formularios = {}
carpeta = "formularios eventos"

for archivo in os.listdir(carpeta):
    ruta = os.path.join(carpeta, archivo)
    try:
        tablas = pd.read_html(ruta)
        archivos_formularios[archivo] = tablas[0]
    except (ValueError, IndexError):
        try:
            df = pd.read_excel(ruta)
            archivos_formularios[archivo] = df
        except Exception:
            pass  # No se pudo leer el archivo, simplemente ignoramos

from unidecode import unidecode  # Importa para eliminar acentos

# Definimos las posibles columnas equivalentes
mapa_columnas = {
    "Nombre": ["Nombre:", "First name:"],
    "Apellidos": ["Apellidos:", "Last name:"],
    "Empresa/Institución": ["Empresa/Institución:", "Company/Institution:"],
    "Miembro de secpho?": ["Miembro de secpho?"],
    "Cargo/Función": ["Cargo/Función:", "Position"],
    "Teléfono": ["Teléfono:", "Telephone:"],
    "Email": ["Email:"]
}

# Diccionario de personas
asistencia_eventos = {}

# Recorrer todos los archivos
for nombre_archivo, df in archivos_formularios.items():

    # Buscar las columnas relevantes en este archivo
    columnas_encontradas = {}
    columnas_archivo = [str(col).strip().lower() for col in df.columns]

    for campo_objetivo, posibles_nombres in mapa_columnas.items():
        for nombre_col in posibles_nombres:
            nombre_col_normalizado = nombre_col.strip().lower()
            for i, col in enumerate(columnas_archivo):
                if nombre_col_normalizado in col:
                    columnas_encontradas[campo_objetivo] = df.columns[i]
                    break
            if campo_objetivo in columnas_encontradas:
                break

    # Recorrer las filas del archivo
    for _, fila in df.iterrows():
        nombre = str(fila.get(columnas_encontradas.get("Nombre", ""), "")).strip()
        apellidos = str(fila.get(columnas_encontradas.get("Apellidos", ""), "")).strip()

        if not nombre or not apellidos:
            continue

        # Crear nombre completo concatenado, sin espacios, en minúsculas y sin acentos
        nombre_completo = unidecode((nombre + apellidos).replace(" ", "").lower())

        if nombre_completo not in asistencia_eventos:
            # Nuevo registro
            asistencia_eventos[nombre_completo] = {
                "Nombre completo": nombre_completo,
                "Eventos asistidos": [nombre_archivo]
            }
        else:
            # Añadir evento si no está ya registrado
            if nombre_archivo not in asistencia_eventos[nombre_completo]["Eventos asistidos"]:
                asistencia_eventos[nombre_completo]["Eventos asistidos"].append(nombre_archivo)

# Convertir a DataFrame
filas = [
    {"Nombre completo": persona["Nombre completo"], "Eventos asistidos": persona["Eventos asistidos"]}
    for persona in asistencia_eventos.values()
]

df_asistencia = pd.DataFrame(filas)
#----------------------------------------------------------------------------------------------------------

# OBTENCIÓN DE LAS TECNOLOGÍAS, ÁMBITOS Y SECTORES DE EVENTOS PASADOS
#----------------------------------------------------------------------------------------------------------
def contar_participacion_por_persona(nombre_persona, df_personas, df_eventos):
    import pandas as pd
    import numpy as np
    import re
    from datetime import datetime, timedelta

    # Verificar si la persona existe
    persona_fila = df_personas[df_personas["Nombre completo"] == nombre_persona]
    if persona_fila.empty:
        # Devolver cuatro DataFrames vacíos con columnas predefinidas
        df_tecnologias = pd.DataFrame(columns=["Tecnología", "Frecuencia_tecnología"])
        df_sectores = pd.DataFrame(columns=["Sector", "Frecuencia_sector"])
        df_ambitos = pd.DataFrame(columns=["Ámbito", "Frecuencia_ámbito"])
        df_eventos_recientes = pd.DataFrame(columns=df_eventos.columns)
        return df_tecnologias, df_sectores, df_ambitos, df_eventos_recientes

    # Obtener eventos asistidos
    eventos_raw = persona_fila["Eventos asistidos"].iloc[0]

    # Limpiar nombres de eventos
    eventos_limpios = []
    for e in eventos_raw:
        if isinstance(e, str):
            e = e.lstrip("'")
            e = re.sub(r"_\d{4}-\d{2}.*", "", e)
            e = e.replace(".xlsx", "")
            e = e.strip()
            eventos_limpios.append(e)

    eventos_limpios = [titulo.replace('_ -', '?:') for titulo in eventos_limpios]
    eventos_limpios = [titulo.replace(' _ ', ' | ') for titulo in eventos_limpios]
    eventos_limpios = [titulo.replace('_', ':') for titulo in eventos_limpios]
    # Filtrar eventos asistidos
    eventos_filtrados = df_eventos[df_eventos["Título"].isin(eventos_limpios)].copy()

    # --- NUEVO: Filtrar eventos de los últimos 365 días ---
    # Asegurarse de que la columna 'Fecha' es tipo datetime
    #if not np.issubdtype(df_eventos["Fecha"].dtype, np.datetime64):
   #     eventos_filtrados["Fecha"] = pd.to_datetime(eventos_filtrados["Fecha"], errors='coerce')

    # Calcular fecha límite
   # fecha_hoy = pd.Timestamp.today()
   # fecha_limite = fecha_hoy - pd.Timedelta(days=365)

    # Filtrar eventos recientes
   # df_eventos_recientes = eventos_filtrados[eventos_filtrados["Fecha"] >= fecha_limite].copy()

    # Inicializar listas
    tecnologias, sectores, ambitos = [], [], []

    # Iterar sobre eventos_filtrados
    for _, fila in eventos_filtrados.iterrows():
        valor_tecnologia = fila.get("Tecnología")
        if isinstance(valor_tecnologia, list):
            tecnologias += [t.strip() for t in valor_tecnologia if isinstance(t, str) and t.strip()]

        valor_sector = fila.get("Sector")
        if isinstance(valor_sector, list):
            sectores += [s.strip() for s in valor_sector if isinstance(s, str) and s.strip()]

        valor_ambito = fila.get("Ámbito")
        if isinstance(valor_ambito, list):
            ambitos += [a.strip() for a in valor_ambito if isinstance(a, str) and a.strip()]

    # Crear dataframes de frecuencia
    df_tecnologias = pd.Series(tecnologias).value_counts().reset_index()
    df_tecnologias.columns = ["Tecnología", "Frecuencia_tecnología"]

    df_sectores = pd.Series(sectores).value_counts().reset_index()
    df_sectores.columns = ["Sector", "Frecuencia_sector"]

    df_ambitos = pd.Series(ambitos).value_counts().reset_index()
    df_ambitos.columns = ["Ámbito", "Frecuencia_ámbito"]

    return df_tecnologias, df_sectores, df_ambitos, eventos_filtrados


#-------------------------------------------------------------------------------------------------------------

# POR SOCIO
#--------------------------------------------------------------------------------------------------------------
from collections import Counter
import pandas as pd
import re

def contar_participacion_por_socio(persona, miembros, df_eventos, df_personas, suscriptores):
    """
    Agrega la participación (tecnologías, sectores, ámbitos, eventos) de
    todas las personas (miembros) que pertenecen al socio indicado en `persona`.
    Devuelve 4 DataFrames: tecnologias, sectores, ambitos, eventos (eventos únicos
    con el nombre de la persona que asistió).
    """

    socio_norm = normalizar_texto(persona.get("Nombre completo", ""))

    miembros_filtrados = miembros[
        miembros["Socio"].apply(normalizar_texto) == socio_norm
    ]

    # Si no hay miembros para ese socio, devolvemos los DataFrames vacíos
    if miembros_filtrados.empty:
        return (
            pd.DataFrame(columns=["Tecnología", "Frecuencia_tecnología"]),
            pd.DataFrame(columns=["Sector", "Frecuencia_sector"]),
            pd.DataFrame(columns=["Ámbito", "Frecuencia_ámbito"]),
            pd.DataFrame(columns=list(df_eventos.columns) + ["Nombre completo"]),
        )

    # Contadores para acumular frecuencias
    contador_tecn = Counter()
    contador_sect = Counter()
    contador_amb  = Counter()

    eventos_list = []  # aquí acumulamos DataFrames de eventos

    # Recorremos cada miembro del socio
    for _, miembro in miembros_filtrados.iterrows():
        nombre_persona = miembro["Nombre completo"]

        # Llamada a la función por persona (debe devolver df_t, df_s, df_a, df_e)
        df_t, df_s, df_a, df_e = contar_participacion_por_persona(
            nombre_persona, df_personas, df_eventos
        )

        # ---- Tecnologías ----
        if not df_t.empty:
            if df_t.shape[1] >= 2:
                key_col, freq_col = df_t.columns[0], df_t.columns[1]
                for _, r in df_t.iterrows():
                    k = r[key_col]
                    if pd.isna(k):
                        continue
                    try:
                        f = int(pd.to_numeric(r[freq_col], errors="coerce") or 0)
                    except Exception:
                        f = 1
                    if f <= 0:
                        f = 1
                    contador_tecn[str(k).strip()] += f
            else:
                key_col = df_t.columns[0]
                for _, r in df_t.iterrows():
                    k = r[key_col]
                    if pd.isna(k):
                        continue
                    contador_tecn[str(k).strip()] += 1

        # ---- Sectores ----
        if not df_s.empty:
            if df_s.shape[1] >= 2:
                key_col, freq_col = df_s.columns[0], df_s.columns[1]
                for _, r in df_s.iterrows():
                    k = r[key_col]
                    if pd.isna(k):
                        continue
                    try:
                        f = int(pd.to_numeric(r[freq_col], errors="coerce") or 0)
                    except Exception:
                        f = 1
                    if f <= 0:
                        f = 1
                    contador_sect[str(k).strip()] += f
            else:
                key_col = df_s.columns[0]
                for _, r in df_s.iterrows():
                    k = r[key_col]
                    if pd.isna(k):
                        continue
                    contador_sect[str(k).strip()] += 1

        # ---- Ámbitos ----
        if not df_a.empty:
            if df_a.shape[1] >= 2:
                key_col, freq_col = df_a.columns[0], df_a.columns[1]
                for _, r in df_a.iterrows():
                    k = r[key_col]
                    if pd.isna(k):
                        continue
                    try:
                        f = int(pd.to_numeric(r[freq_col], errors="coerce") or 0)
                    except Exception:
                        f = 1
                    if f <= 0:
                        f = 1
                    contador_amb[str(k).strip()] += f
            else:
                key_col = df_a.columns[0]
                for _, r in df_a.iterrows():
                    k = r[key_col]
                    if pd.isna(k):
                        continue
                    contador_amb[str(k).strip()] += 1

        # ---- Eventos ----
        if not df_e.empty:
            df_e = df_e.copy()
            df_e["Nombre completo"] = nombre_persona
            eventos_list.append(df_e)

    # ----- Procesar suscriptores -----
    suscriptores_filtrados = suscriptores[
        suscriptores["Socio"].apply(normalizar_texto) == socio_norm
    ]

    for _, sub in suscriptores_filtrados.iterrows():
        nombre_sub = sub["Nombre completo"]

        # Buscar la fila correspondiente en df_personas
        fila_persona = df_personas[df_personas["Nombre completo"] == nombre_sub]

        if not fila_persona.empty:
            eventos_raw = fila_persona.iloc[0].get("Eventos asistidos", "")

            eventos_separados = []
            if isinstance(eventos_raw, str):
                eventos_separados = [e.strip() for e in eventos_raw.split(",") if e.strip()]
            elif isinstance(eventos_raw, list):
                eventos_separados = [str(e).strip() for e in eventos_raw if str(e).strip()]

            # --- LIMPIAR NOMBRES DE EVENTOS ---
            eventos_limpios = []
            for e in eventos_separados:
                if isinstance(e, str):
                    e = e.lstrip("'")
                    e = re.sub(r"_\d{4}-\d{2}.*", "", e)
                    e = e.replace(".xlsx", "")
                    e = e.strip()
                    eventos_limpios.append(e)

            eventos_limpios = [titulo.replace('_ -', '?:') for titulo in eventos_limpios]
            eventos_limpios = [titulo.replace(' _ ', ' | ') for titulo in eventos_limpios]
            eventos_limpios = [titulo.replace('_', ':') for titulo in eventos_limpios]

            # Crear DataFrame de eventos sin fecha
            if eventos_limpios:
                df_e_sub = pd.DataFrame({
                    "Título": eventos_limpios,
                    "Nombre completo": nombre_sub
                })
                eventos_list.append(df_e_sub)

    # ---- Construir DataFrames finales a partir de contadores ----
    df_tecnologias = (pd.DataFrame(contador_tecn.items(), columns=["Tecnología", "Frecuencia_tecnología"])
                      .sort_values("Frecuencia_tecnología", ascending=False)
                      .reset_index(drop=True)) if contador_tecn else pd.DataFrame(columns=["Tecnología", "Frecuencia_tecnología"])

    df_sectores = (pd.DataFrame(contador_sect.items(), columns=["Sector", "Frecuencia_sector"])
                   .sort_values("Frecuencia_sector", ascending=False)
                   .reset_index(drop=True)) if contador_sect else pd.DataFrame(columns=["Sector", "Frecuencia_sector"])

    df_ambitos = (pd.DataFrame(contador_amb.items(), columns=["Ámbito", "Frecuencia_ámbito"])
                  .sort_values("Frecuencia_ámbito", ascending=False)
                  .reset_index(drop=True)) if contador_amb else pd.DataFrame(columns=["Ámbito", "Frecuencia_ámbito"])

    # ---- Concatenar eventos y eliminar duplicados, agrupar por persona ----
    if eventos_list:
        df_eventos_all = pd.concat(eventos_list, ignore_index=True)
        if "Título" in df_eventos_all.columns:
            df_eventos_all = df_eventos_all.drop_duplicates(subset=["Título", "Nombre completo"]).reset_index(drop=True)
            df_eventos_all = df_eventos_all.groupby("Nombre completo", as_index=False).agg({
                "Título": lambda x: ", ".join(sorted(x))
            })
    else:
        df_eventos_all = pd.DataFrame(columns=list(df_eventos.columns) + ["Nombre completo"])

    # ---- Mapear Nombre completo a Nombre + Apellidos ----
    def reemplazar_nombre_completo(df_eventos_all, miembros, suscriptores):
        mapping = {}

        for _, row in miembros.iterrows():
            nombre_completo = row.get("Nombre completo")
            nombre_apellidos = f"{row.get('Nombre', '').strip()} {row.get('Apellidos', '').strip()}".strip()
            if nombre_completo and nombre_apellidos:
                mapping[nombre_completo] = nombre_apellidos

        for _, row in suscriptores.iterrows():
            nombre_completo = row.get("Nombre completo")
            nombre_apellidos = f"{row.get('Nombre', '').strip()} {row.get('Apellidos', '').strip()}".strip()
            if nombre_completo and nombre_apellidos:
                mapping[nombre_completo] = nombre_apellidos

        df_eventos_all["Nombre completo"] = df_eventos_all["Nombre completo"].map(lambda x: mapping.get(x, x))
        return df_eventos_all

    df_eventos_all = reemplazar_nombre_completo(df_eventos_all, miembros, suscriptores)

    return df_tecnologias, df_sectores, df_ambitos, df_eventos_all

#------------------------------------------------------------------------------------------------------------
def reformatear_eventos_por_evento(df_eventos_all, eventos_pasados):
    """
    Convierte un DataFrame con una fila por persona y eventos en una fila por evento,
    con la lista de personas asistentes y la fecha del evento (buscada en eventos_pasados).
    
    Entradas:
    - df_eventos_all: DataFrame con columnas ["Nombre completo", "Título"] (cadena de eventos separados por coma)
    - eventos_pasados: DataFrame con al menos ["Título", "Fecha"]

    Salida:
    - df_eventos_por_evento: DataFrame con columnas ["Título", "Fecha", "Asistentes"]
    """

    if df_eventos_all.empty:
        return pd.DataFrame(columns=["Título", "Fecha", "Asistentes"])

    # Convertir a lista de filas: una por (persona, evento)
    filas = []
    for _, row in df_eventos_all.iterrows():
        persona = row["Nombre completo"]
        titulos = row["Título"]
        if pd.isna(titulos):
            continue
        eventos = [e.strip() for e in titulos.split(",") if e.strip()]
        for evento in eventos:
            filas.append({"Título": evento, "Nombre completo": persona})

    df_expandido = pd.DataFrame(filas)

    # Agrupar por Título y juntar asistentes
    df_eventos_por_evento = (
        df_expandido
        .groupby("Título", as_index=False)
        .agg({"Nombre completo": lambda x: ", ".join(sorted(set(x)))})
        .rename(columns={"Nombre completo": "Asistentes"})
    )

    # Buscar fecha del evento desde `eventos_pasados`
    if "Fecha" in eventos_pasados.columns:
        eventos_pasados_reducido = eventos_pasados[["Título", "Fecha"]].drop_duplicates()
        df_eventos_por_evento = df_eventos_por_evento.merge(eventos_pasados_reducido, on="Título", how="left")
    else:
        df_eventos_por_evento["Fecha"] = None

    # Ordenar por fecha si está disponible
    if df_eventos_por_evento["Fecha"].notna().any():
        df_eventos_por_evento = df_eventos_por_evento.sort_values("Fecha").reset_index(drop=True)

    return df_eventos_por_evento

# FUNCIÓN PARA RECOMENDAR EVENTOS PERSONA
#-------------------------------------------------------------------------------------------------------------
def recomendar_eventos_con_historial(nombre_persona, df_personas, df_eventos_futuros,
                                      freq_tecn, freq_sect, freq_amb):
    from unidecode import unidecode
    import numpy as np
    import pandas as pd

    # Ponderaciones
    PESOS = {
        "tec": 29.0,
        "sec": 15.0,
        "amb": 12.0,
        "prov": 22.0,
        "tec_hist": 10.0,
        "sec_hist": 7.0,
        "amb_hist": 5.0,
        "online": 22.0
    }

    def limpiar_str(s):
        return unidecode(str(s).strip().lower())

    def normalizar_lista(val):
        """Convierte string o lista en lista de strings limpios"""
        if isinstance(val, list):
            return [limpiar_str(v) for v in val if isinstance(v, str) and v.strip()]
        elif isinstance(val, str):
            return [limpiar_str(x) for x in val.split(",") if x.strip()]
        else:
            return []

    # Obtener la persona
    persona_fila = df_personas[df_personas["Nombre completo"] == nombre_persona]
    if persona_fila.empty:
        raise ValueError(f"No se encontró la persona con nombre completo '{nombre_persona}'")
    persona = persona_fila.iloc[0]

    intereses_tecn = normalizar_lista(persona.get("Tecnologías", ""))
    intereses_sect = normalizar_lista(persona.get("Sectores", ""))
    intereses_amb = normalizar_lista(persona.get("Ámbitos", ""))
    provincia_persona = limpiar_str(persona.get("Provincia prof.", ""))

    recomendaciones = []

    for _, evento in df_eventos_futuros.iterrows():
        # Listas del evento
        tecnos_evento = normalizar_lista(evento.get("Tecnología", []))
        sectos_evento = normalizar_lista(evento.get("Sector", []))
        ambits_evento = normalizar_lista(evento.get("Ámbito", []))
        provincia_evento = limpiar_str(evento.get("Provincia", ""))
        ubicacion_evento = limpiar_str(evento.get("Ubicación", ""))

        score = 0.0

        # --- Coincidencia de tecnologías ---
        if "todas" in tecnos_evento:
            score += PESOS["tec"]
        elif tecnos_evento:
            inter = set(tecnos_evento) & set(intereses_tecn)
            proporción = len(inter) / len(tecnos_evento)
            score += proporción * PESOS["tec"]

        # --- Coincidencia de sectores ---
        if "todos" in sectos_evento:
            score += PESOS["sec"]
        elif sectos_evento:
            inter = set(sectos_evento) & set(intereses_sect)
            proporción = len(inter) / len(sectos_evento)
            score += proporción * PESOS["sec"]

        # --- Coincidencia de ámbitos ---
        if ambits_evento:
            inter = set(ambits_evento) & set(intereses_amb)
            proporción = len(inter) / len(ambits_evento)
            score += proporción * PESOS["amb"]

        # --- Coincidencia de provincia ---
        if provincia_persona and provincia_persona == provincia_evento:
            score += PESOS["prov"]

        # --- Historial de tecnologías ---
        if tecnos_evento:
            inter = set(tecnos_evento) & set(freq_tecn.keys())
            proporción = len(inter) / len(tecnos_evento)
            score += proporción * PESOS["tec_hist"]

        # --- Historial de sectores ---
        if sectos_evento:
            inter = set(sectos_evento) & set(freq_sect.keys())
            proporción = len(inter) / len(sectos_evento)
            score += proporción * PESOS["sec_hist"]

        # --- Historial de ámbitos ---
        if ambits_evento:
            inter = set(ambits_evento) & set(freq_amb.keys())
            proporción = len(inter) / len(ambits_evento)
            score += proporción * PESOS["amb_hist"]

        # --- Bonus por ser online ---
        if "online" in ubicacion_evento:
            score += PESOS["online"]

        # No superar 100%
        score = min(score, 100.0)

        recomendaciones.append({
            "Evento": evento.get("Título", "Sin título"),
            "Score": f"{round(score, 1)}%"
        })

    # Crear y ordenar el DataFrame
    df_recomendaciones = pd.DataFrame(recomendaciones)
    df_recomendaciones["Score_num"] = df_recomendaciones["Score"].str.rstrip('%').astype(float)
    df_recomendaciones = df_recomendaciones.sort_values(by="Score_num", ascending=False).drop(columns=["Score_num"]).reset_index(drop=True)

    return df_recomendaciones
#-----------------------------------------------------------------------------------------------------------------------

# FUNCIÓN PARA RECOMENDAR EVENTOS SOCIOS
#-------------------------------------------------------------------------------------------------------------------
def recomendar_eventos_con_historial_socio(persona, df_eventos_futuros,
                                      freq_tecn, freq_sect, freq_amb):
    from unidecode import unidecode
    import pandas as pd

    # Ponderaciones
    PESOS = {
        "tec": 29.0,
        "sec": 15.0,
        "amb": 12.0,
        "prov": 22.0,
        "tec_hist": 10.0,
        "sec_hist": 7.0,
        "amb_hist": 5.0,
        "online": 22.0
    }

    def limpiar_str(s):
        return unidecode(str(s).strip().lower())

    def normalizar_lista(val):
        """Convierte string o lista en lista de strings limpios"""
        if isinstance(val, list):
            return [limpiar_str(v) for v in val if isinstance(v, str) and v.strip()]
        elif isinstance(val, str):
            return [limpiar_str(x) for x in val.split(",") if x.strip()]
        else:
            return []

    # 🔹 Ahora tomamos los datos directamente de la fila persona
    intereses_tecn = normalizar_lista(persona.get("Tecnologías", ""))
    intereses_sect = normalizar_lista(persona.get("Sectores", ""))
    intereses_amb = normalizar_lista(persona.get("Ámbitos", ""))
    provincia_persona = limpiar_str(persona.get("Provincia prof.", ""))

    recomendaciones = []

    for _, evento in df_eventos_futuros.iterrows():
        # Listas del evento
        tecnos_evento = normalizar_lista(evento.get("Tecnología", []))
        sectos_evento = normalizar_lista(evento.get("Sector", []))
        ambits_evento = normalizar_lista(evento.get("Ámbito", []))
        provincia_evento = limpiar_str(evento.get("Provincia", ""))
        ubicacion_evento = limpiar_str(evento.get("Ubicación", ""))

        score = 0.0

        # --- Coincidencia de tecnologías ---
        if "todas" in tecnos_evento:
            score += PESOS["tec"]
        elif tecnos_evento:
            inter = set(tecnos_evento) & set(intereses_tecn)
            proporcion = len(inter) / len(tecnos_evento)
            score += proporcion * PESOS["tec"]

        # --- Coincidencia de sectores ---
        if "todos" in sectos_evento:
            score += PESOS["sec"]
        elif sectos_evento:
            inter = set(sectos_evento) & set(intereses_sect)
            proporcion = len(inter) / len(sectos_evento)
            score += proporcion * PESOS["sec"]

        # --- Coincidencia de ámbitos ---
        if ambits_evento:
            inter = set(ambits_evento) & set(intereses_amb)
            proporcion = len(inter) / len(ambits_evento)
            score += proporcion * PESOS["amb"]

        # --- Coincidencia de provincia ---
        if provincia_persona and provincia_persona == provincia_evento:
            score += PESOS["prov"]

        # --- Historial de tecnologías ---
        if tecnos_evento:
            inter = set(tecnos_evento) & set(freq_tecn.keys())
            proporcion = len(inter) / len(tecnos_evento)
            score += proporcion * PESOS["tec_hist"]

        # --- Historial de sectores ---
        if sectos_evento:
            inter = set(sectos_evento) & set(freq_sect.keys())
            proporcion = len(inter) / len(sectos_evento)
            score += proporcion * PESOS["sec_hist"]

        # --- Historial de ámbitos ---
        if ambits_evento:
            inter = set(ambits_evento) & set(freq_amb.keys())
            proporcion = len(inter) / len(ambits_evento)
            score += proporcion * PESOS["amb_hist"]

        # --- Bonus por ser online ---
        if "online" in ubicacion_evento:
            score += PESOS["online"]

        # No superar 100%
        score = min(score, 100.0)

        recomendaciones.append({
            "Evento": evento.get("Título", "Sin título"),
            "Score": f"{round(score, 1)}%"
        })

    # Crear y ordenar el DataFrame
    df_recomendaciones = pd.DataFrame(recomendaciones)
    if not df_recomendaciones.empty:
        df_recomendaciones["Score_num"] = df_recomendaciones["Score"].str.rstrip('%').astype(float)
        df_recomendaciones = (
            df_recomendaciones
            .sort_values(by="Score_num", ascending=False)
            .drop(columns=["Score_num"])
            .reset_index(drop=True)
        )

    return df_recomendaciones
#------------------------------------------------------------------------------------------------------------------

# OBTENCIÓN DE LOS DATOS DE RETOS TECNOLÓGICOS
#-----------------------------------------------------------------------------------------------------------------------
url = "https://secpho.org/wp-json/reports/v1/retos?auth=c9J2mL7vT8sW4xAfB3eR6zNpQ1HdUgVKtrXa"

headers = {
    "Accept": "application/json",
    "User-Agent": "Mozilla/5.0"
}

response = requests.get(url, headers=headers)

if response.status_code == 200:
    data = response.json()
    retos_tecnológicos = pd.DataFrame(data)
else:
    retos_tecnológicos = pd.DataFrame()  # Devuelve un DataFrame vacío si hay error
retos_tecnológicos=retos_tecnológicos.T

def reemplazar_guion_emisor(valor):
    if isinstance(valor, str):
        return valor.replace(' &#8211;', '-')
    elif isinstance(valor, list):
        return [str(item).replace(' &#8211;', '-') for item in valor]
    else:
        return valor  # Para NaN u otros tipos
def reemplazar_guion_emisorr(valor):
    if isinstance(valor, str):
        return valor.replace('&#038;', '&')
    elif isinstance(valor, list):
        return [str(item).replace('&#038;', '&') for item in valor]
    else:
        return valor  # Para NaN u otros tipos

retos_tecnológicos['Entidad emisora'] = retos_tecnológicos['Entidad emisora'].apply(reemplazar_guion_emisor)
retos_tecnológicos['Entidades que aplican'] = retos_tecnológicos['Entidades que aplican'].apply(reemplazar_guion_emisor)
retos_tecnológicos['Entidad emisora'] = retos_tecnológicos['Entidad emisora'].apply(reemplazar_guion_emisorr)
retos_tecnológicos['Entidades que aplican'] = retos_tecnológicos['Entidades que aplican'].apply(reemplazar_guion_emisorr)
# Asegúrate de que la columna 'Fecha' sea de tipo datetime
retos_tecnológicos["Fecha cierre"] = pd.to_datetime(retos_tecnológicos["Fecha cierre"], errors="coerce")

# Obtener la fecha actual
hoy = pd.Timestamp.today().normalize()  # Normaliza para ignorar la hora

# Filtrar los retos cuya fecha es posterior a hoy
retos_futuros = retos_tecnológicos[retos_tecnológicos["Fecha cierre"] > hoy].copy()
#------------------------------------------------------------------------------------------------------------------

# FUNCIÓN PARA RECOMENDAR RETOS
#-----------------------------------------------------------------------------------------------------------------
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import pandas as pd
import numpy as np

# Cargar modelo de embeddings
modelo = SentenceTransformer("all-MiniLM-L6-v2")

def recomendar_retos_por_perfil(nombre_persona, df_personas, df_retos, top_n=5):
    # Verificar que la persona exista
    persona_fila = df_personas[df_personas["Nombre completo"] == nombre_persona]
    if persona_fila.empty:
        raise ValueError(f"No se encontró la persona '{nombre_persona}'")

    persona = persona_fila.iloc[0]

    # --- 1. Crear texto representativo del perfil ---
    def limpiar_y_concatenar(val):
        if isinstance(val, list):
            return ", ".join(val)
        elif isinstance(val, str):
            return val
        else:
            return ""

    perfil_texto = " ".join([
        limpiar_y_concatenar(persona.get("Subtecnologías", "")),
        limpiar_y_concatenar(persona.get("Subsectores", "")),
        limpiar_y_concatenar(persona.get("Ámbitos", "")),
    ])

    # Embedding del perfil
    embedding_perfil = modelo.encode(perfil_texto)

    # Sectores de la persona (lista limpia en minúsculas)
    sectores_persona = set([s.strip().lower() for s in persona.get("Sectores", [])]) if isinstance(persona.get("Sectores", []), list) else set()

    # --- 2. Preparar texto y sectores de cada reto ---
    textos_retos = []
    sectores_retos = []
    for _, reto in df_retos.iterrows():
        texto = " ".join([
            str(reto.get("Título", "")),
            str(reto.get("Descripción", "")),
            limpiar_y_concatenar(reto.get("Sector/es", "")),
        ])
        textos_retos.append(texto)

        sectores = reto.get("Sector/es", [])
        if isinstance(sectores, str):  # por si viniera como texto
            sectores = [s.strip() for s in sectores.split(",") if s.strip()]
        sectores = [s.lower() for s in sectores]
        sectores_retos.append(set(sectores))

    # --- 3. Calcular embeddings de los retos ---
    embeddings_retos = modelo.encode(textos_retos)

    # --- 4. Calcular similitud coseno ---
    similitudes = cosine_similarity([embedding_perfil], embeddings_retos)[0]

    # --- 5. Calcular coincidencia de sectores ---
    coincidencias = []
    for s_reto in sectores_retos:
        if not s_reto:  # si el reto no tiene sectores definidos
            coincidencias.append(0.0)
        else:
            interseccion = len(s_reto & sectores_persona)
            coincidencias.append(interseccion / len(s_reto))  # proporción de solapamiento

    # --- 6. Score final (50% sectores + 50% similitud coseno) ---
    scores_finales = 0.5 * similitudes + 0.5 * np.array(coincidencias)

    # --- 7. Crear DataFrame de resultados ---
    df_resultados = df_retos.copy()
    df_resultados["Similitud coseno"] = similitudes
    df_resultados["Coincidencia sectores"] = coincidencias
    df_resultados["ScoreFinal"] = scores_finales

    df_resultados = df_resultados.sort_values(by="ScoreFinal", ascending=False).head(top_n).reset_index(drop=True)

    return df_resultados
#--------------------------------------------------------------------------------------------------------------------

# RECOMENDAR RETOS PARA SOCIOS
#-------------------------------------------------------------------------------------------------------------------
def recomendar_retos_por_perfil_socio(persona, df_retos, top_n=5):
    from sklearn.metrics.pairwise import cosine_similarity

    # --- 1. Crear texto representativo del perfil ---
    def limpiar_y_concatenar(val):
        if isinstance(val, list):
            return ", ".join(val)
        elif isinstance(val, str):
            return val
        else:
            return ""

    perfil_texto = " ".join([
        limpiar_y_concatenar(persona.get("Subtecnologías", "")),
        limpiar_y_concatenar(persona.get("Subsectores", "")),
        limpiar_y_concatenar(persona.get("Ámbitos", ""))
    ])

    # Embedding del perfil
    embedding_perfil = modelo.encode(perfil_texto)

    # --- 2. Preparar texto de cada reto ---
    textos_retos = []
    for _, reto in df_retos.iterrows():
        texto = " ".join([
            str(reto.get("Título", "")),
            str(reto.get("Descripción", "")),
            limpiar_y_concatenar(reto.get("Sector/es", ""))
        ])
        textos_retos.append(texto)

    # --- 3. Calcular embeddings de los retos ---
    embeddings_retos = modelo.encode(textos_retos)

    # --- 4. Calcular similitud coseno ---
    similitudes = cosine_similarity([embedding_perfil], embeddings_retos)[0]

    # --- 5. Calcular coincidencia de sectores (50%) ---
    sectores_persona = set(map(str.lower, persona.get("Sectores", []))) if isinstance(persona.get("Sectores", []), list) else set()

    coincidencias = []
    for _, reto in df_retos.iterrows():
        sectores_reto = reto.get("Sector/es", [])
        if isinstance(sectores_reto, str):
            sectores_reto = [sectores_reto]
        sectores_reto = set(map(str.lower, sectores_reto))

        if len(sectores_reto) > 0:
            interseccion = len(sectores_reto & sectores_persona)
            coincidencia = interseccion / len(sectores_reto)  # proporción cubierta
        else:
            coincidencia = 0.0

        coincidencias.append(coincidencia)

    # --- 6. Score combinado ---
    df_resultados = df_retos.copy()
    df_resultados["Similitud"] = similitudes
    df_resultados["CoincidenciaSectores"] = coincidencias
    df_resultados["ScoreFinal"] = 0.5 * df_resultados["Similitud"] + 0.5 * df_resultados["CoincidenciaSectores"]

    # --- 7. Ordenar y devolver ---
    df_resultados = (
        df_resultados.sort_values(by="ScoreFinal", ascending=False)
        .head(top_n)
        .reset_index(drop=True)
    )

    return df_resultados
#------------------------------------------------------------------------------------------------------------------

# RETOS DEL ÚLTIMO AÑO
#--------------------------------------------------------------------------------------------------------------------
import pandas as pd

# Asegurarse de que la columna sea datetime
retos_tecnológicos["Fecha envío"] = pd.to_datetime(retos_tecnológicos["Fecha envío"], errors="coerce")

# Fecha actual
hoy = pd.Timestamp.now()

# Filtrar: solo los que ya pasaron
retos_pasados = retos_tecnológicos[retos_tecnológicos["Fecha envío"] < hoy]
#--------------------------------------------------------------------------------------------------------------------

# CARGA DE DATOS DE PROYECTOS
#-------------------------------------------------------------------------------------------------------------------
import pandas as pd

# Cargar únicamente la hoja llamada "Retos"
proyectos = pd.read_excel("Datos de Proyectos.xlsx", sheet_name="Proyectos")

df_sector_terms = pd.read_excel(
    "Datos de Proyectos.xlsx",
    sheet_name="sectores",
    header=1   # fila 1 → segunda línea (la primera es 0)
)

df_tecnologia_terms = pd.read_excel(
    "Datos de Proyectos.xlsx",
    sheet_name="Tecnologias",
    header=2   # fila 1 → segunda línea (la primera es 0)
)

df_ambitos_terms = pd.read_excel(
    "Datos de Proyectos.xlsx",
    sheet_name="ámbitos",
    header=2   # fila 1 → segunda línea (la primera es 0)
)
df_entidades_terms = pd.read_excel(
    "Datos de Proyectos.xlsx",
    sheet_name="entidades",
    header=0   # fila 1 → segunda línea (la primera es 0)
)

# Crear el diccionario de mapeo desde df_sector_terms
mapa_sector = dict(zip(df_sector_terms["ID"].astype(str), df_sector_terms["Nombre"]))

# Función para reemplazar múltiples IDs por nombres
def traducir_sectores(celda):
    if pd.isna(celda):
        return ""
    ids = [id_.strip() for id_ in str(celda).split(",")]
    nombres = [mapa_sector.get(id_, f"[ID no encontrado: {id_}]") for id_ in ids]
    return ", ".join(nombres)

# Aplicar al DataFrame df
proyectos["Sectores"] = proyectos["Sectores"].apply(traducir_sectores)

mapa_tecnologia = dict(zip(df_tecnologia_terms["ID"].astype(str), df_tecnologia_terms["Nombre"]))

# Función para reemplazar múltiples IDs por nombres
def traducir_tecnologia(celda):
    if pd.isna(celda):
        return ""
    # Si es 'proyecto no tecnológico', dejar tal cual
    if str(celda).strip().lower() == "proyecto no tecnológico":
        return celda
    # Traducir IDs
    ids = [id_.strip() for id_ in str(celda).split(",")]
    nombres = [mapa_tecnologia.get(id_, f"[ID no encontrado: {id_}]") for id_ in ids]
    return ", ".join(nombres)

# Aplicar al DataFrame df
proyectos["Tecnologías"] = proyectos["Tecnologías"].apply(traducir_tecnologia)

mapa_ambito = dict(zip(df_ambitos_terms["ID"].astype(str), df_ambitos_terms["Nombre"]))

def traducir_ambito(celda):
    if pd.isna(celda):
        return ""
    # Si es 'proyecto no tecnológico', dejar tal cual
    if str(celda).strip().lower() == "todos":
        return celda
    # Traducir IDs
    ids = [id_.strip() for id_ in str(celda).split(",")]
    nombres = [mapa_ambito.get(id_, f"[ID no encontrado: {id_}]") for id_ in ids]
    return ", ".join(nombres)

# Aplicar al DataFrame df
proyectos["Ámbitos"] = proyectos["Ámbitos"].apply(traducir_ambito)

mapa_entidades = dict(zip(df_entidades_terms["ID"].astype(str), df_entidades_terms["Título"]))

# Función para reemplazar múltiples IDs por nombres
def traducir_entidades(celda):
    if pd.isna(celda):
        return ""
    ids = [id_.strip() for id_ in str(celda).split(",")]
    nombres = [mapa_entidades.get(id_, f"[ID no encontrado: {id_}]") for id_ in ids]
    return ", ".join(nombres)

# Aplicar al DataFrame df
proyectos["Partners"] = proyectos["Partners"].apply(traducir_entidades)

def reemplazar_guion_emisor(valor):
    if isinstance(valor, str):
        return valor.replace(' &#8211;', '-')
    elif isinstance(valor, list):
        return [str(item).replace(' &#8211;', '-') for item in valor]
    else:
        return valor  # Para NaN u otros tipos
def reemplazar_guion_emisorr(valor):
    if isinstance(valor, str):
        return valor.replace('&#038;', '&')
    elif isinstance(valor, list):
        return [str(item).replace('&#038;', '&') for item in valor]
    else:
        return valor  # Para NaN u otros tipos

proyectos['Partners'] = proyectos['Partners'].apply(reemplazar_guion_emisor)
proyectos['Partners'] = proyectos['Partners'].apply(reemplazar_guion_emisorr)
#------------------------------------------------------------------------------------------------------------------

# BUSQUEDA DE PROYECTOS DE UNA EMPRESA
#----------------------------------------------------------------------------------------------------------------
def proyectos_de_socio(df_proyectos, socio):
    """
    Devuelve un DataFrame con los proyectos en los que participa un socio.

    Parámetros:
        df_proyectos (pd.DataFrame): DataFrame con una columna 'Partners'.
        socio (str): Nombre del socio a buscar.

    Retorna:
        pd.DataFrame: Subconjunto de proyectos donde aparece el socio.
    """
    socio = socio.strip().lower()

    # Filtramos proyectos donde la columna Partners contenga el socio
    mask = df_proyectos['Partners'].fillna("").str.lower().str.split(",").apply(
        lambda socios: any(socio == s.strip() for s in socios)
    )

    return df_proyectos[mask].copy()
#-----------------------------------------------------------------------------------------------------------------

# BÚSQUEDA DE PROYECTOS CON COINCIDENCIA EN TECNOLOGÍA O SECTOR
#-----------------------------------------------------------------------------------------------------------------
def proyectos_relacionados(proyectos, tecnologias_persona, sectores_persona):
    """
    Busca proyectos relacionados con una persona según tecnologías y sectores.
    - Primero: proyectos que coinciden en al menos 1 tecnología Y 1 sector.
    - Si no hay: proyectos que coinciden en tecnología O en sector.
    Devuelve como máximo los 5 proyectos más recientes (según 'Final').
    """

    # --- Normalizar entrada ---
    def normalizar_lista(x):
        if pd.isna(x):
            return []
        if isinstance(x, str):  
            return [i.strip().lower() for i in x.split(",") if i.strip()]
        elif isinstance(x, list):  
            return [str(i).strip().lower() for i in x if pd.notna(i)]
        else:
            return [str(x).strip().lower()]

    tecnologias_persona = normalizar_lista(tecnologias_persona)
    sectores_persona = normalizar_lista(sectores_persona)

    proyectos = proyectos.copy()

    # Asegurar que "Final" es datetime
    if " Final" in proyectos.columns:
        proyectos[" Final"] = pd.to_datetime(proyectos[" Final"], errors="coerce")

    proyectos_filtrados = []

    # --- Evaluar coincidencias ---
    for _, fila in proyectos.iterrows():
        tec = normalizar_lista(fila.get("Tecnologías", ""))
        sec = normalizar_lista(fila.get("Sectores", ""))

        match_tecnologias = set(tec) & set(tecnologias_persona)
        match_sectores = set(sec) & set(sectores_persona)

        fila_copy = fila.copy()
        fila_copy["Tecnología de coincidencia"] = ", ".join(match_tecnologias)
        fila_copy["Sector de coincidencia"] = ", ".join(match_sectores)

        # Guardamos siempre, luego filtramos según condición
        fila_copy["_match_tecnologias"] = bool(match_tecnologias)
        fila_copy["_match_sectores"] = bool(match_sectores)

        proyectos_filtrados.append(fila_copy)

    df_filtrado = pd.DataFrame(proyectos_filtrados)

    if df_filtrado.empty:
        return df_filtrado

    # --- Paso 1: buscar coincidencias en ambos ---
    df_and = df_filtrado[(df_filtrado["_match_tecnologias"]) & (df_filtrado["_match_sectores"])]

    if not df_and.empty:
        return df_and.sort_values(" Final", ascending=False).head(5).drop(columns=["_match_tecnologias", "_match_sectores"])

    # --- Paso 2: si no hay dobles coincidencias, buscar simples ---
    df_or = df_filtrado[(df_filtrado["_match_tecnologias"]) | (df_filtrado["_match_sectores"])]

    return df_or.sort_values(" Final", ascending=False).head(5).drop(columns=["_match_tecnologias", "_match_sectores"])
#--------------------------------------------------------------------------------------------------------------------

# GENERACIÓN DEL INFORME DE PERSONA
#--------------------------------------------------------------------------------------------------------------------
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.shared import RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from io import BytesIO

def generar_informe_persona(nombre_persona):
    doc = Document()

    # Cambiar estilo Normal
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'DM Sans'
    font_normal.size = Pt(11)
    style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'DM Sans')
    paragraph_format = style_normal.paragraph_format
    paragraph_format.line_spacing = 1.0
    # Espacio antes del párrafo (en puntos)
    paragraph_format.space_before = 2

    # Espacio después del párrafo (en puntos)
    paragraph_format.space_after = 2


    section = doc.sections[0]
    header = section.header

    p = header.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Añadir tabulador derecho en la posición deseada (por ejemplo, 16 cm)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

    # Añadir "Pag. "
    run = p.add_run("Pag. ")

    # Campo dinámico número de página
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "1"  # Texto temporal, Word lo reemplaza

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)

    font = run.font
    font.size = Pt(9)

    # Añadir tabulador para separar texto de imagen
    p.add_run("\t")

    # Añadir imagen a la derecha
    run_img = p.add_run()
    run_img.add_picture('logo1.png', width=Inches(0.75))
    
    

    # Crear estilo personalizado solo si no existe
    if 'CustomTitle' not in [s.name for s in doc.styles]:
        style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'DM Sans'
        font.size = Pt(18)
        font.bold = True

        # Forzar rFonts
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'DM Sans')
        rFonts.set(qn('w:hAnsi'), 'DM Sans')
        rFonts.set(qn('w:eastAsia'), 'DM Sans')
        rFonts.set(qn('w:cs'), 'DM Sans')
        style.element.rPr.insert(0, rFonts)

    if 'CustomTitle1' not in [s.name for s in doc.styles]:
        style = doc.styles.add_style('CustomTitle1', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'DM Sans'
        font.size = Pt(16)
        font.bold = True

        # Forzar rFonts
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'DM Sans')
        rFonts.set(qn('w:hAnsi'), 'DM Sans')
        rFonts.set(qn('w:eastAsia'), 'DM Sans')
        rFonts.set(qn('w:cs'), 'DM Sans')
        style.element.rPr.insert(0, rFonts)

    # Crear estilo personalizado solo si no existe
    if 'CustomTitle2' not in [s.name for s in doc.styles]:
        style = doc.styles.add_style('CustomTitle2', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'DM Sans'
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = RGBColor(255, 255, 255)

        # Forzar rFonts
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'DM Sans')
        rFonts.set(qn('w:hAnsi'), 'DM Sans')
        rFonts.set(qn('w:eastAsia'), 'DM Sans')
        rFonts.set(qn('w:cs'), 'DM Sans')
        style.element.rPr.insert(0, rFonts)

    # Añadir título centrado con estilo personalizado
    titulo = doc.add_paragraph("Informe de Valor y Oportunidades", style='CustomTitle')
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    

    # Buscar la persona en el dataframe
    persona_fila = miembros.loc[miembros["Nombre completo"] == nombre_persona]
    if persona_fila.empty:
        doc.add_paragraph("⚠️ No se encontró información sobre esta persona.")
        return doc  # Devolvemos el doc aunque esté incompleto
    
    persona = persona_fila.iloc[0]  # Extraer la fila

    
    def set_paragraph_background(paragraph, color_hex):
    # Añadir sombreado a todo el párrafo
        p = paragraph._element
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)  # Ej: 'FFFF00' para amarillo
        p.get_or_add_pPr().append(shading)

    # Añadir resumen de información personal
    p = doc.add_paragraph("Resumen de la persona", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(p, "FF7E27")
    
    # b
    b = doc.add_paragraph()
    b.add_run("Nombre: ").bold = True
    b.add_run(f"{persona.get('Nombre', 'N/D')} {persona.get('Apellidos', 'N/D')}")

    # c
    c = doc.add_paragraph()
    c.add_run("Empresa: ").bold = True
    c.add_run(f"{persona.get('Socio', 'N/D')}")

    # d
    p = doc.add_paragraph()
    p.add_run("Función: ").bold = True
    p.add_run(f"{persona.get('Función', 'N/D')}")

    # e
    e = doc.add_paragraph()
    e.add_run("Cargo: ").bold = True
    e.add_run(f"{persona.get('Cargo', 'N/D')}")

    # f
    f = doc.add_paragraph()
    f.add_run("Provincia profesional: ").bold = True
    f.add_run(f"{persona.get('Provincia prof.', 'N/D')}")

    # g
    g = doc.add_paragraph()
    g.add_run("Tecnologías: ").bold = True
    g.add_run(f"{persona.get('Tecnologías', 'N/D')}")

    # h
    h = doc.add_paragraph()
    h.add_run("Sectores: ").bold = True
    h.add_run(f"{persona.get('Sectores', 'N/D')}")

    # i
    ambitos = persona.get("Ámbitos", [])
    if isinstance(ambitos, list):
        ambitos_str = ", ".join(ambitos)
    elif isinstance(ambitos, str):
        ambitos_str = ambitos.strip("[]").replace("'", "").replace('"', '')
    else:
        ambitos_str = str(ambitos)

    i = doc.add_paragraph()
    i.add_run("Ámbitos: ").bold = True
    i.add_run(ambitos_str)
    doc.add_paragraph("")

    

    tecnologias_eventos_pasados, sectores_eventos_pasados, ambitos_eventos_pasados, eventos_asistidos_persona = contar_participacion_por_persona(nombre_persona, df_asistencia, eventos_pasados)
    freq_tecn = dict(zip(tecnologias_eventos_pasados["Tecnología"].str.lower(), tecnologias_eventos_pasados["Frecuencia_tecnología"]))
    freq_sect = dict(zip(sectores_eventos_pasados["Sector"].str.lower(), sectores_eventos_pasados["Frecuencia_sector"]))
    freq_amb = dict(zip(ambitos_eventos_pasados["Ámbito"].str.lower(), ambitos_eventos_pasados["Frecuencia_ámbito"]))
    recomendaciones_eventos = recomendar_eventos_con_historial(nombre_persona, miembros, próximos_eventos, freq_tecn, freq_sect, freq_amb)

    # Filtrar eventos con puntuación positiva
    # Convertir Score a número antes de filtrar
    recomendaciones_eventos["Score_num"] = recomendaciones_eventos["Score"].str.rstrip('%').astype(float)

    # Filtrar eventos con puntuación positiva
    recomendaciones_positivas = recomendaciones_eventos[recomendaciones_eventos["Score_num"] > 0].head(3)
    
    # Añadir título centrado con estilo personalizado
    titulo = doc.add_paragraph("Eventos", style='CustomTitle1')
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Añadir sección de recomendaciones al informe
    j=doc.add_paragraph("Recomendación de Eventos", style='CustomTitle2')
    j.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(j, "4570F7")

    if not recomendaciones_positivas.empty:
        for idx, fila in recomendaciones_positivas.iterrows():
            evento_info = próximos_eventos[próximos_eventos["Título"] == fila["Evento"]].iloc[0]

            # Título y Score en negrita
            a = doc.add_paragraph(style='List Bullet')
            a.add_run(evento_info['Título']).bold = True
            a.add_run(f" (Score: {fila['Score_num']:.2f}%)").bold = True


            # Tecnologías
            tecnos = evento_info.get("Tecnología", [])
            if isinstance(tecnos, list):
                tecnos_str = ", ".join(tecnos)
            elif isinstance(tecnos, str):
                tecnos_str = tecnos.strip("[]").replace("'", "").replace('"', '')
            else:
                tecnos_str = str(tecnos)
            p = doc.add_paragraph()
            p.add_run("- Tecnologías del evento: ").bold = True
            p.add_run(tecnos_str)

            # Sector
            sectos = evento_info.get("Sector", [])
            if isinstance(sectos, list):
                sectos_str = ", ".join(sectos)
            elif isinstance(sectos, str):
                sectos_str = sectos.strip("[]").replace("'", "").replace('"', '')
            else:
                sectos_str = str(sectos)
            p = doc.add_paragraph()
            p.add_run("- Sectores del evento: ").bold = True
            p.add_run(sectos_str)

            # Ámbito
            ambis = evento_info.get("Ámbito", [])
            if isinstance(ambis, list):
                ambis_str = ", ".join(ambis)
            elif isinstance(ambis, str):
                ambis_str = ambis.strip("[]").replace("'", "").replace('"', '')
            else:
                ambis_str = str(ambis)
            p = doc.add_paragraph()
            p.add_run("- Ámbitos del evento: ").bold = True
            p.add_run(ambis_str)

            # Ubicación
            p = doc.add_paragraph()
            p.add_run("- Ubicación: ").bold = True
            p.add_run(str(evento_info.get('Ubicación', 'N/D')))

            # Provincia (solo si la ubicación no es 'online')
            ubicacion = str(evento_info.get('Ubicación', '')).strip().lower()
            if ubicacion != 'online':
                p = doc.add_paragraph()
                p.add_run("- Provincia donde se celebrará: ").bold = True
                p.add_run(str(evento_info.get('Provincia', 'N/D')))

            
            from datetime import datetime
            
            import sys

            

            fecha_evento = evento_info.get('Fecha', 'N/D')

            def formatear_fecha_es(fecha):
                if not fecha or pd.isna(fecha):
                    return "Fecha desconocida"
                if isinstance(fecha, datetime):
                    fmt = '%#d de %B de %Y' if sys.platform.startswith('win') else '%-d de %B de %Y'
                    return fecha.strftime(fmt)
                else:
                    try:
                        fecha_obj = datetime.fromisoformat(str(fecha))
                        fmt = '%#d de %B de %Y' if sys.platform.startswith('win') else '%-d de %B de %Y'
                        return fecha_obj.strftime(fmt)
                    except ValueError:
                        try:
                            fecha_obj = datetime.strptime(str(fecha), "%Y-%m-%d %H:%M:%S")
                            fmt = '%#d de %B de %Y' if sys.platform.startswith('win') else '%-d de %B de %Y'
                            return fecha_obj.strftime(fmt)
                        except ValueError:
                            return str(fecha).split()[0]

            fecha_str = formatear_fecha_es(fecha_evento)

            p = doc.add_paragraph()
            p.add_run("- Fecha: ").bold = True
            p.add_run(fecha_str)


            doc.add_paragraph("")  # Línea en blanco entre eventos
    else:
        doc.add_paragraph("No se encontraron eventos relevantes para recomendar en este momento.")
        doc.add_paragraph("")

    from datetime import datetime, timedelta

    # Añadir sección de eventos asistidos en el último año
    p=doc.add_paragraph(f"Eventos a los que ha asistido {persona.get('Nombre', 'N/D')} {persona.get('Apellidos', 'N/D')}", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(p, "4570F7")

    
    if not eventos_asistidos_persona.empty:
        for _, evento in eventos_asistidos_persona.iterrows():
            try:
                fecha_dt = pd.to_datetime(evento["Fecha"], errors="coerce")
                fecha_str = fecha_dt.strftime("%d/%m/%Y") if pd.notna(fecha_dt) else "Fecha desconocida"
            except Exception:
                fecha_str = "Fecha desconocida"
            doc.add_paragraph(f"{evento['Título']} ({fecha_str})", style='ListBullet')
    else:
        doc.add_paragraph("Esta persona no ha asistido a ningún evento organizado por SECPhO.")
        

    doc.add_paragraph("")
    # Añadir título centrado con estilo personalizado
    titulo = doc.add_paragraph("Retos Tecnológicos", style='CustomTitle1')
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Recomendación de retos
    p=doc.add_paragraph("Recomendación de Retos Tecnológicos Activos", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(p, "F25830")

    recomendaciones_retos = recomendar_retos_por_perfil(nombre_persona, miembros, retos_futuros, top_n=10)
    recomendaciones_retos["Score_num"] = recomendaciones_retos["ScoreFinal"] * 100  # numérico
    recomendaciones_retos["Similitud %"] = (recomendaciones_retos["ScoreFinal"] * 100).round(2).astype(str) + "%"  # visual

    # Filtrar eventos con puntuación positiva
    recomendaciones_positivas_retos = recomendaciones_retos[recomendaciones_retos["Score_num"] > 0].head(3)

    if not recomendaciones_positivas_retos.empty:
        for idx, fila in recomendaciones_positivas_retos.iterrows():
            # Obtener entidad emisora del reto
            reto_info = retos_futuros[retos_futuros["Título"] == fila["Título"]].iloc[0]
            empresa_persona = str(persona.get('Socio', 'N/D')).strip()

            # --- Comprobar si la empresa de la persona coincide con entidad emisora ---
            entidad_emisora = reto_info.get("Entidad emisora", "")
            entidades_emisoras = []

            if pd.notna(entidad_emisora):
                if isinstance(entidad_emisora, list):
                    entidades_emisoras = [str(e).strip() for e in entidad_emisora]
                elif isinstance(entidad_emisora, str):
                    entidades_emisoras = [e.strip() for e in entidad_emisora.split(",")]
                else:
                    entidades_emisoras = [str(entidad_emisora).strip()]

            if empresa_persona in entidades_emisoras:
                continue  # Saltar si la empresa de la persona es la emisora del reto

            # --- Comprobar si la empresa de la persona ya ha aplicado ---
            entidades_aplican = reto_info.get("Entidades que aplican", [])
            entidades_aplicantes = []

            if pd.notna(entidades_aplican):
                if isinstance(entidades_aplican, list):
                    entidades_aplicantes = [str(e).strip() for e in entidades_aplican]
                elif isinstance(entidades_aplican, str):
                    entidades_aplicantes = [e.strip() for e in entidades_aplican.split(",")]
                else:
                    entidades_aplicantes = [str(entidades_aplican).strip()]

            if empresa_persona in entidades_aplicantes:
                continue  # Saltar si la empresa de la persona ya ha aplicado

            # --- Si pasó ambas comprobaciones, se muestra el reto ---

            # Título y Score en negrita
            a = doc.add_paragraph(style='List Bullet')
            a.add_run(reto_info['Título']).bold = True
            a.add_run(f" (Score de similitud: {fila['Score_num']:.2f}%)").bold = True

            # Descripción
            descripcion = reto_info.get("Descripción", "")
            if not descripcion or pd.isna(descripcion) or str(descripcion).strip() == "":
                descripcion = "Descripción no disponible."

            p = doc.add_paragraph()
            p.add_run("- Descripción: ").bold = True
            p.add_run(descripcion)

            # Sector
            sector = reto_info.get("Sector/es", [])
            if isinstance(sector, list):
                sector_str = ", ".join(sector)
            elif isinstance(sector, str):
                sector_str = sector.strip("[]").replace("'", "").replace('"', '')
            else:
                sector_str = str(sector)

            p = doc.add_paragraph()
            p.add_run("- Sector al que pertenece el reto: ").bold = True
            p.add_run(sector_str)

            # Entidad emisora
            p = doc.add_paragraph()
            p.add_run("- Entidad emisora del reto: ").bold = True
            p.add_run(str(reto_info.get('Entidad emisora', 'N/D')))

            # Entidades que aplican
            if not entidades_aplican:
                entidades_texto = "Ninguna"
            else:
                entidades_texto = str(entidades_aplican)

            p = doc.add_paragraph()
            p.add_run("- Entidades que ya han aplicado: ").bold = True
            p.add_run(entidades_texto)

            # Fecha de cierre
            fecha_cierre = reto_info.get('Fecha cierre', 'N/D')
            fecha_cierre_str = formatear_fecha_es(fecha_cierre)

            p = doc.add_paragraph()
            p.add_run("- Fecha de cierre: ").bold = True
            p.add_run(fecha_cierre_str)

            doc.add_paragraph("")  # Espacio al final
    else:
        doc.add_paragraph("No se han encontrado retos tecnológicos para recomendar a esta persona.")
        doc.add_paragraph("")
            
    # Añadir sección de retos en el último año
    empresa_persona = str(persona.get('Socio', 'N/D')).strip()
    p=doc.add_paragraph(f"Retos tecnológicos emitidos por {empresa_persona} que han sido mediados por SECPhO", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(p, "F25830")

    # Lista para guardar retos de la empresa
    retos_de_su_empresa = []

    for _, fila in retos_pasados.iterrows():
        entidad_emisora = fila["Entidad emisora"]

        entidades = []
        if pd.notna(entidad_emisora):
            if isinstance(entidad_emisora, list):
                entidades = [str(e).strip() for e in entidad_emisora]
            elif isinstance(entidad_emisora, str):
                entidades = [e.strip() for e in entidad_emisora.split(",")]
            else:
                entidades = [str(entidad_emisora).strip()]
        
        if empresa_persona in entidades:
            retos_de_su_empresa.append(fila)

    # Mostrar resultados
    if not retos_de_su_empresa:
        doc.add_paragraph("Su empresa nunca ha emitido un reto tecnológico.")
        doc.add_paragraph("")
    else:
        for fila in retos_de_su_empresa:
            p = doc.add_paragraph(style='ListBullet')
            p.add_run(f"{fila['Título']}").bold = True


            # Descripción
            descripcion = fila.get("Descripción", "")
            if not descripcion or pd.isna(descripcion) or str(descripcion).strip() == "":
                descripcion = "Descripción no disponible."
            p = doc.add_paragraph()
            p.add_run("- Descripción: ").bold = True
            p.add_run(descripcion)


            # Sector
            sector = fila["Sector/es"]
            if isinstance(sector, list):
                sector_str = ", ".join(sector)
            elif isinstance(sector, str):
                sector_str = sector.strip("[]").replace("'", "").replace('"', '')
            else:
                sector_str = str(sector)

            p = doc.add_paragraph()
            p.add_run("- Sector al que pertenece el reto: ").bold = True
            p.add_run(sector_str)

            # Entidades que aplican
            entidades_aplican = fila.get('Entidades que aplican', None)
            entidades_texto = "Ninguna" if not entidades_aplican else str(entidades_aplican)

            p = doc.add_paragraph()
            p.add_run("- Entidades que ya han aplicado: ").bold = True
            p.add_run(entidades_texto)

            


            # Fecha de cierre
            fecha_cierre = fila['Fecha cierre']
            fecha_cierre_str = formatear_fecha_es(fecha_cierre)
            p = doc.add_paragraph()
            p.add_run("- Fecha de cierre: ").bold = True
            p.add_run(fecha_cierre_str)
            doc.add_paragraph("")


    # Añadir sección de retos en el último año
    empresa_persona = str(persona.get('Socio', 'N/D')).strip()
    p=doc.add_paragraph(f"Retos tecnológicos en los que {empresa_persona} ha aplicado", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(p, "F25830")

    # Lista para guardar retos de la empresa
    retos_de_su_empresa = []

    for _, fila in retos_pasados.iterrows():
        entidades_aplicantes = fila["Entidades que aplican"]

        entidades = []
        if pd.notna(entidades_aplicantes):
            if isinstance(entidades_aplicantes, list):
                entidades = [str(e).strip() for e in entidades_aplicantes]
            elif isinstance(entidades_aplicantes, str):
                entidades = [e.strip() for e in entidades_aplicantes.split(",")]
            else:
                entidades = [str(entidades_aplicantes).strip()]
        
        if empresa_persona in entidades:
            retos_de_su_empresa.append(fila)

    # Mostrar resultados
    if not retos_de_su_empresa:
        doc.add_paragraph("Su empresa nunca ha aplicado a un reto tecnológico.")
        doc.add_paragraph("")
    else:
        for fila in retos_de_su_empresa:
            p = doc.add_paragraph(style='ListBullet')
            p.add_run(f"{fila['Título']}").bold = True


            # Descripción
            descripcion = fila.get("Descripción", "")
            if not descripcion or pd.isna(descripcion) or str(descripcion).strip() == "":
                descripcion = "Descripción no disponible."
            p = doc.add_paragraph()
            p.add_run("- Descripción: ").bold = True
            p.add_run(descripcion)


            # Sector
            sector = fila["Sector/es"]
            if isinstance(sector, list):
                sector_str = ", ".join(sector)
            elif isinstance(sector, str):
                sector_str = sector.strip("[]").replace("'", "").replace('"', '')
            else:
                sector_str = str(sector)

            p = doc.add_paragraph()
            p.add_run("- Sector al que pertenece el reto: ").bold = True
            p.add_run(sector_str)

            # Entidad emisora
            emisor = fila["Entidad emisora"]
            if isinstance(emisor, list):
                emisor_str = ", ".join(emisor)
            elif isinstance(emisor, str):
                emisor_str = emisor.strip("[]").replace("'", "").replace('"', '')
            else:
                emisor_str = str(emisor)

            p = doc.add_paragraph()
            p.add_run("- Entidad emisora del reto: ").bold = True
            p.add_run(emisor_str)

            # Entidades que aplican
            entidades_aplican = fila.get('Entidades que aplican', None)
            entidades_texto = "Ninguna" if not entidades_aplican else str(entidades_aplican)

            p = doc.add_paragraph()
            p.add_run("- Entidades que ya han aplicado: ").bold = True
            p.add_run(entidades_texto)

            # Fecha de cierre
            fecha_cierre = fila['Fecha cierre']
            fecha_cierre_str = formatear_fecha_es(fecha_cierre)
            p = doc.add_paragraph()
            p.add_run("- Fecha de cierre: ").bold = True
            p.add_run(fecha_cierre_str) 
            doc.add_paragraph("")   


    # Añadir título centrado con estilo personalizado
    titulo = doc.add_paragraph("Proyectos", style='CustomTitle1')
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Añadir sección de proyectos de la empresa
    empresa_persona = str(persona.get('Socio', 'N/D')).strip()
    p=doc.add_paragraph(f"Proyectos en los que {empresa_persona} ha sido Partner", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(p, "4C49F2")

    proyectos_filtrados = proyectos_de_socio(proyectos, empresa_persona) 
    if proyectos_filtrados.empty:
        doc.add_paragraph("Su empresa no ha sido partner de ningún proyecto.")
        doc.add_paragraph("")
    else:
        for _, fila in proyectos_filtrados.iterrows():
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{fila['Título de la Idea/Proyecto']}").bold = True



            # Fecha de inicio
            fecha_inicio = fila[' Inicio']
            fecha_inicio_str = formatear_fecha_es(fecha_inicio)
            p = doc.add_paragraph()
            p.add_run("- Fecha de inicio del proyecto: ").bold = True
            p.add_run(fecha_inicio_str)

            # Fecha de cierre
            fecha_final = fila[' Final']
            fecha_final_str = formatear_fecha_es(fecha_final)
            p = doc.add_paragraph()
            p.add_run("- Fecha de finalización del proyecto: ").bold = True
            p.add_run(fecha_final_str)

            # partners
            part = fila["Partners"]
            if isinstance(part, list):
                part_str = ", ".join(part)
            elif isinstance(part, str):
                part_str = part.strip("[]").replace("'", "").replace('"', '')
            else:
                part_str = str(part)
            if part_str == "":
                part_str = 'No definidos'
            p = doc.add_paragraph()
            p.add_run("- Partners del proyecto: ").bold = True
            p.add_run(part_str)

            # Origen fondos
            p = doc.add_paragraph()
            p.add_run("- Origen de los fondos: ").bold = True
            p.add_run(str(fila.get('Origen fondos ', 'N/D')))

            # Programa financiación
            programa = fila.get("Programa financiacion ", "")
            if not programa or pd.isna(programa) or str(programa).strip() == "":
                programa = "No consta"
            p = doc.add_paragraph()
            p.add_run("- Programa de financiación: ").bold = True
            p.add_run(programa)

            # Presupuesto total
            presu = fila.get("Presupuesto total (€)", "")
            if pd.notna(presu):
                p = doc.add_paragraph()
                p.add_run("- Presupuest total: ").bold = True
                p.add_run(f"{presu:,.2f} €")  
            else:
                p = doc.add_paragraph()
                p.add_run("- Presupuest total: ").bold = True
                p.add_run("No consta")




            # tecnologias
            tecno = fila["Tecnologías"]
            if isinstance(tecno, list):
                tecno_str = ", ".join(tecno)
            elif isinstance(tecno, str):
                tecno_str = tecno.strip("[]").replace("'", "").replace('"', '')
            else:
                tecno_str = str(tecno)
            if tecno_str == "":
                tecno_str = 'No definidas'
            p = doc.add_paragraph()
            p.add_run("- Tecnologías del proyecto: ").bold = True
            p.add_run(tecno_str)

            # Sector
            sector = fila["Sectores"]
            if isinstance(sector, list):
                sector_str = ", ".join(sector)
            elif isinstance(sector, str):
                sector_str = sector.strip("[]").replace("'", "").replace('"', '')
            else:
                sector_str = str(sector)
            if sector_str == "":
                sector_str = 'No definido'
            p = doc.add_paragraph()
            p.add_run("- Sectores del proyecto: ").bold = True
            p.add_run(sector_str)

            # ambitos
            ambi = fila["Ámbitos"]
            if isinstance(ambi, list):
                ambi_str = ", ".join(ambi)
            elif isinstance(ambi, str):
                ambi_str = ambi.strip("[]").replace("'", "").replace('"', '')
            else:
                ambi_str = str(ambi)
            if ambi_str == "":
                ambi_str = 'No definido'
            p = doc.add_paragraph()
            p.add_run("- Ámbitos del proyecto: ").bold = True
            p.add_run(ambi_str)
            doc.add_paragraph("")

    tecnologias_persona = persona.get('Tecnologías')
    sectores_persona=persona.get('Sectores')

    proyectos_rel = proyectos_relacionados(proyectos, tecnologias_persona, sectores_persona)

    p=doc.add_paragraph(f"Información sobre proyectos pasados coordinados por SECPhO relacionados con las tecnologías y sectores de interés de {persona.get('Nombre', 'N/D')} {persona.get('Apellidos', 'N/D')}", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(p, "4C49F2") 

    if proyectos_rel.empty:
        doc.add_paragraph("No se ha realizado ningún proyecto de interés para esta persona.")
        doc.add_paragraph("")
    else:
        def esta_vacio(x):
            if x is None:
                return True
            if isinstance(x, float) and np.isnan(x):
                return True
            if isinstance(x, str):
                s = x.strip()
                return (s == "" or s.lower() in {"nan", "n/d", "nd", "none"})
            if isinstance(x, (list, tuple, set)):
                return len(x) == 0
            return False

        def fmt_valor(v):
            if isinstance(v, (list, tuple, set)):
                return ", ".join(str(x) for x in v)
            return str(v)

        for _, fila in proyectos_rel.iterrows():
            tec = fila.get('Tecnología de coincidencia')
            sec = fila.get('Sector de coincidencia')

            partes = []
            if not esta_vacio(tec):
                partes.append(f"Tecnología: {fmt_valor(tec)}")
            if not esta_vacio(sec):
                partes.append(f"Sector: {fmt_valor(sec)}")

            # Si no hay nada que mostrar, saltar
            if not partes:
                continue

            p = doc.add_paragraph(" || ".join(partes), style='CustomTitle2')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_paragraph_background(p, "8C8AFF")

    
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{fila['Título de la Idea/Proyecto']}").bold = True



            # Fecha de inicio
            fecha_inicio = fila[' Inicio']
            fecha_inicio_str = formatear_fecha_es(fecha_inicio)
            p = doc.add_paragraph()
            p.add_run("- Fecha de inicio del proyecto: ").bold = True
            p.add_run(fecha_inicio_str)

            # Fecha de cierre
            fecha_final = fila[' Final']
            fecha_final_str = formatear_fecha_es(fecha_final)
            p = doc.add_paragraph()
            p.add_run("- Fecha de finalización del proyecto: ").bold = True
            p.add_run(fecha_final_str)

            # partners
            part = fila["Partners"]
            if isinstance(part, list):
                part_str = ", ".join(part)
            elif isinstance(part, str):
                part_str = part.strip("[]").replace("'", "").replace('"', '')
            else:
                part_str = str(part)
            if part_str == "":
                part_str = 'No definidos'
            p = doc.add_paragraph()
            p.add_run("- Partners del proyecto: ").bold = True
            p.add_run(part_str)

            # Origen fondos
            p = doc.add_paragraph()
            p.add_run("- Origen de los fondos: ").bold = True
            p.add_run(str(fila.get('Origen fondos ', 'N/D')))

            # Programa financiación
            programa = fila.get("Programa financiacion ", "")
            if not programa or pd.isna(programa) or str(programa).strip() == "":
                programa = "No consta"
            p = doc.add_paragraph()
            p.add_run("- Programa de financiación: ").bold = True
            p.add_run(programa)

            # Presupuesto total
            presu = fila.get("Presupuesto total (€)", "")
            if pd.notna(presu):
                p = doc.add_paragraph()
                p.add_run("- Presupuest total: ").bold = True
                p.add_run(f"{presu:,.2f} €")  
            else:
                p = doc.add_paragraph()
                p.add_run("- Presupuest total: ").bold = True
                p.add_run("No consta")



            # tecnologias
            tecno = fila["Tecnologías"]
            if isinstance(tecno, list):
                tecno_str = ", ".join(tecno)
            elif isinstance(tecno, str):
                tecno_str = tecno.strip("[]").replace("'", "").replace('"', '')
            else:
                tecno_str = str(tecno)
            if tecno_str == "":
                tecno_str = 'No definidas'
            p = doc.add_paragraph()
            p.add_run("- Tecnologías del proyecto: ").bold = True
            p.add_run(tecno_str)

            # Sector
            sector = fila["Sectores"]
            if isinstance(sector, list):
                sector_str = ", ".join(sector)
            elif isinstance(sector, str):
                sector_str = sector.strip("[]").replace("'", "").replace('"', '')
            else:
                sector_str = str(sector)
            if sector_str == "":
                sector_str = 'No definido'
            p = doc.add_paragraph()
            p.add_run("- Sectores del proyecto: ").bold = True
            p.add_run(sector_str)

            # ambitos
            ambi = fila["Ámbitos"]
            if isinstance(ambi, list):
                ambi_str = ", ".join(ambi)
            elif isinstance(ambi, str):
                ambi_str = ambi.strip("[]").replace("'", "").replace('"', '')
            else:
                ambi_str = str(ambi)
            if ambi_str == "":
                ambi_str = 'No definido'
            p = doc.add_paragraph()
            p.add_run("- Ámbitos del proyecto: ").bold = True
            p.add_run(ambi_str)
            doc.add_paragraph("")




    # Guardar en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# GENERACIÓN INFORME SOCIO
#-----------------------------------------------------------------------------------------------------------
def generar_informe_socio(nombre_persona):
    doc = Document()

    # Cambiar estilo Normal
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'DM Sans'
    font_normal.size = Pt(11)
    style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'DM Sans')
    paragraph_format = style_normal.paragraph_format
    paragraph_format.line_spacing = 1.0
    paragraph_format.space_before = 2
    paragraph_format.space_after = 2
    
    section = doc.sections[0]
    

    # 👇 Esta línea es clave
    section.different_first_page_header_footer = True
    # Crear estilo personalizado solo si no existe
    if 'CustomTitle' not in [s.name for s in doc.styles]:
        style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'DM Sans'
        font.size = Pt(18)
        font.bold = True

        # Forzar rFonts
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'DM Sans')
        rFonts.set(qn('w:hAnsi'), 'DM Sans')
        rFonts.set(qn('w:eastAsia'), 'DM Sans')
        rFonts.set(qn('w:cs'), 'DM Sans')
        style.element.rPr.insert(0, rFonts)
    # === Crear tabla con 1 fila y 2 columnas ===
    table = doc.add_table(rows=1, cols=2)
    table.allow_autofit = True
    table.autofit = True
    
    # Ajustar anchos de las columnas (por ejemplo, 10 cm y 8 cm)
    table.columns[0].width = Inches(4)  # columna de imagen (~10.16 cm)
    table.columns[1].width = Inches(3)  # columna de texto (~7.62 cm)
    
    # === Celda izquierda: imagen ===
    cell_img = table.cell(0, 0)
    paragraph_img = cell_img.paragraphs[0]
    run_img = paragraph_img.add_run()
    run_img.add_picture('imagen_portada.jpg', width=Inches(4))  # Ajusta tamaño si necesario
    
    # === Celda derecha: título ===
    cell_title = table.cell(0, 1)
    paragraph_title = cell_title.paragraphs[0]
    paragraph_title.style = 'CustomTitle'
    paragraph_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    persona_fila = socios.loc[socios["Nombre completo"] == nombre_persona]
    persona = persona_fila.iloc[0]  # Extraer la fila
    run_title = paragraph_title.add_run(f"Informe de Valor y Oportunidades para {persona.get('Socio', 'N/D')}")
    doc.add_page_break()
    
    
    # ========================
    # ENCABEZADO (solo imagen a la derecha)
    # ========================
    header = section.header
    p_header = header.add_paragraph()
    p_header.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run_img = p_header.add_run()
    run_img.add_picture('logo1.png', width=Inches(1.00))  # Puedes ajustar el tamaño
    
    # ========================
    # PIE DE PÁGINA (número de página a la derecha)
    # ========================
    footer = section.footer
    p_footer = footer.add_paragraph()
    p_footer.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
    # Campo "Pag. {PAGE}"
    run = p_footer.add_run("Pag. ")
    
    # Campo dinámico de número de página
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "1"
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)
    
    font = run.font
    font.size = Pt(9)
    
    if 'IndexTitle' not in [s.name for s in doc.styles]:
        style = doc.styles.add_style('IndexTitle', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'DM Sans'
        font.size = Pt(18)
        font.bold = False  # Sin negrita
    
        # Forzar rFonts para DM Sans
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'DM Sans')
        rFonts.set(qn('w:hAnsi'), 'DM Sans')
        rFonts.set(qn('w:eastAsia'), 'DM Sans')
        rFonts.set(qn('w:cs'), 'DM Sans')
        style.element.rPr.insert(0, rFonts)
    
    p_index_title = doc.add_paragraph('Índice', style='CustomTitle')
    nombre_entidad = persona.get('Socio', 'N/D')


    # Lista manual (puedes ajustar tabulación o numeración como prefieras)
    indice_items = [
        "1. Introducción",
        f"2. Resumen de datos de {nombre_entidad}",
        "3. Contactos recomendados",
        "4. Eventos y actividades",
        "5. Retos tecnológicos",
        "6. Proyectos"
    ]
    
    for item in indice_items:
        p = doc.add_paragraph(item, style='Normal')
        p.paragraph_format.space_before = Pt(6)
    doc.add_page_break()
   

    if 'CustomTitle1' not in [s.name for s in doc.styles]:
        style = doc.styles.add_style('CustomTitle1', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'DM Sans'
        font.size = Pt(16)
        font.bold = True

        # Forzar rFonts
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'DM Sans')
        rFonts.set(qn('w:hAnsi'), 'DM Sans')
        rFonts.set(qn('w:eastAsia'), 'DM Sans')
        rFonts.set(qn('w:cs'), 'DM Sans')
        style.element.rPr.insert(0, rFonts)

    # Crear estilo personalizado solo si no existe
    if 'CustomTitle2' not in [s.name for s in doc.styles]:
        style = doc.styles.add_style('CustomTitle2', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'DM Sans'
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = RGBColor(255, 255, 255)

        # Forzar rFonts
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'DM Sans')
        rFonts.set(qn('w:hAnsi'), 'DM Sans')
        rFonts.set(qn('w:eastAsia'), 'DM Sans')
        rFonts.set(qn('w:cs'), 'DM Sans')
        style.element.rPr.insert(0, rFonts)

    # Añadir título centrado con estilo personalizado
    titulo = doc.add_paragraph("1. Introducción", style='CustomTitle')
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    

    # Buscar la persona en el dataframe
    persona_fila = socios.loc[socios["Nombre completo"] == nombre_persona]
    if persona_fila.empty:
        doc.add_paragraph("⚠️ No se encontró información sobre este socio.")
        return doc  # Devolvemos el doc aunque esté incompleto
    else:
        persona = persona_fila.iloc[0]  # Extraer la fila
        nombre_entidad = persona.get('Socio', 'N/D')
        doc.add_paragraph(f"Este informe tiene como objetivo poner en valor la participación de {nombre_entidad} en el ecosistema SECPHO, destacando tanto su trayectoria como las oportunidades que pueden surgir dentro del ecosistema.")
        doc.add_paragraph("")
        doc.add_paragraph("Se recogen recomendaciones de potenciales contactos, así como eventos o actividades, retos tecnológicos y proyectos, tanto pasados como futuros, alineados con sus áreas de interés.")
    
    persona = persona_fila.iloc[0]  # Extraer la fila

    import unicodedata
    import re
    import pandas as pd

    def normalizar_texto(texto):
        if pd.isna(texto):
                return ""
        texto = str(texto).lower()
        texto = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("utf-8")
        texto = re.sub(r"[^a-z0-9]", "", texto)  # quita espacios y símbolos
        return texto

    def agregar_info_a_socio(persona, miembros):
        # Normalizar nombre del socio de la fila actual
        socio_norm = normalizar_texto(persona.get("Nombre completo", ""))

        # Filtrar miembros que pertenezcan a ese socio
        miembros_filtrados = miembros[
            miembros["Socio"].apply(normalizar_texto) == socio_norm
        ]

        def procesar_columna(columna):
            """Extrae valores únicos de una columna de tipo string (coma-separada)"""
            valores = []
            for valor in miembros_filtrados[columna].dropna():
                for v in str(valor).split(","):
                    valores.append(v.strip())
            # Eliminar duplicados preservando orden
            return ", ".join(dict.fromkeys(valores)) if valores else "N/D"

        def procesar_lista(columna):
            """Extrae valores únicos de columnas que ya son listas"""
            valores = []
            for lista in miembros_filtrados[columna].dropna():
                if isinstance(lista, list):  # ya es lista
                    for v in lista:
                        valores.append(str(v).strip())
                else:  # por si viene mal cargado como string
                    for v in str(lista).split(","):
                        valores.append(v.strip())
            return ", ".join(dict.fromkeys(valores)) if valores else "N/D"

        # Procesar cada columna y asignarla a persona
        persona["Tecnologías"] = procesar_columna("Tecnologías")
        persona["Subtecnologías"] = procesar_columna("Subtecnologías")
        persona["Subsectores"] = procesar_columna("Subsectores")
        persona["Sectores"] = procesar_columna("Sectores")
        persona["Ámbitos"] = procesar_lista("Ámbitos")

        return persona

    persona = agregar_info_a_socio(persona, miembros)

    
    def set_paragraph_background(paragraph, color_hex):
    # Añadir sombreado a todo el párrafo
        p = paragraph._element
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color_hex)  # Ej: 'FFFF00' para amarillo
        p.get_or_add_pPr().append(shading)

# Añadir título centrado con estilo personalizado
    titulo = doc.add_paragraph(f"2. Resumen de datos de {persona.get('Socio', 'N/D')}", style='CustomTitle')
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"Este apartado recoge la información clave sobre el perfil de {persona.get('Socio', 'N/D')}, destacando sus ámbitos de actuación, tecnologías y sectores estratégicos")
# Añadir resumen de información personal
    p = doc.add_paragraph("Ficha de socio", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_paragraph_background(p, "FF7E27")
    
    # b
    b = doc.add_paragraph()
    b.add_run("Socio: ").bold = True
    b.add_run(f"{persona.get('Socio', 'N/D')}")

    # c
    c = doc.add_paragraph()
    c.add_run("Tipo de empresa: ").bold = True
    c.add_run(f"{persona.get('Tipo de empresa', 'N/D')}")

    # d
    p = doc.add_paragraph()
    p.add_run("Tipo de socio: ").bold = True
    p.add_run(f"{persona.get('Tipo de socio', 'N/D')}")

    # e
    e = doc.add_paragraph()
    e.add_run("Público o privado: ").bold = True
    e.add_run(f"{persona.get('Pub./Priv.', 'N/D')}")

    # f
    f = doc.add_paragraph()
    f.add_run("Cadena de valor: ").bold = True
    f.add_run(f"{persona.get('Cadena de valor', 'N/D')}")

    # g
    # Obtener el nombre completo del socio actual
    nombre_completo = persona.get("Nombre completo", None)

    # Buscar la provincia en contacto_socios
    if nombre_completo:
        provincia = contacto_socios.loc[
            contacto_socios["Nombre completo"] == nombre_completo, "F_Provincia"
        ]
        if not provincia.empty:
            provincia_valor = provincia.iloc[0]
        else:
            provincia_valor = "N/D"
    else:
        provincia_valor = "N/D"

    persona["Provincia prof."] = provincia_valor

    # Escribir en el documento
    g = doc.add_paragraph()
    g.add_run("Provincia: ").bold = True
    g.add_run(str(provincia_valor))
    
    # h
    h = doc.add_paragraph()
    h.add_run("Tecnologías: ").bold = True
    h.add_run(f"{persona.get('Tecnologías', 'N/D')}")

    # h
    h = doc.add_paragraph()
    h.add_run("Sectores: ").bold = True
    h.add_run(f"{persona.get('Sectores', 'N/D')}")

    
    i = doc.add_paragraph()
    i.add_run("Ámbitos: ").bold = True
    i.add_run(f"{persona.get('Ámbitos', 'N/D')}")
    doc.add_paragraph("")

    titulo = doc.add_paragraph("3. Contactos Recomendados", style='CustomTitle')
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"En este apartado os sugerimos contactos de expertos del ecosistema SECPHO que pueden resultar de interés para miembros del equipo de {persona.get('Socio', 'N/D')}. El objetivo es facilitaros conexiones y, si os interesa, reuniones personalizadas con personas afines a vuestras capacidades, tecnologías clave y áreas de interés, fomentando así nuevas oportunidades de colaboración.")
# Añadir resumen de información personal
    p = doc.add_paragraph("Propuesta de contactos recomendados", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_paragraph_background(p, "F25830")
    doc.add_page_break()

    doc.add_page_break()

    df_tecno_socios, df_secto_socios, df_ambi_socios, df_evento_socios = contar_participacion_por_socio(persona, miembros, eventos_pasados, df_asistencia, suscriptores)
    df_eventos_por_evento = reformatear_eventos_por_evento(df_evento_socios, eventos_pasados)
    freq_tecn = dict(zip(df_tecno_socios["Tecnología"].str.lower(), df_tecno_socios["Frecuencia_tecnología"]))
    freq_sect = dict(zip(df_secto_socios["Sector"].str.lower(), df_secto_socios["Frecuencia_sector"]))
    freq_amb = dict(zip(df_ambi_socios["Ámbito"].str.lower(), df_ambi_socios["Frecuencia_ámbito"]))
    recomendaciones_eventos = recomendar_eventos_con_historial_socio(persona, próximos_eventos, freq_tecn, freq_sect, freq_amb)

    # Filtrar eventos con puntuación positiva
    # Convertir Score a número antes de filtrar
    recomendaciones_eventos["Score_num"] = recomendaciones_eventos["Score"].str.rstrip('%').astype(float)

    # Filtrar eventos con puntuación positiva
    recomendaciones_positivas = recomendaciones_eventos[recomendaciones_eventos["Score_num"] > 0].head(3)

    titulo = doc.add_paragraph("4. Eventos y actividades", style='CustomTitle')
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"Se presentan las recomendaciones de próximos eventos relevantes para {persona.get('Socio', 'N/D')}, así como el histórico de participación de sus miembros en actividades organizadas o dinamizadas por SECPHO.")
    # Añadir título centrado con estilo personalizado
    
    # Añadir sección de recomendaciones al informe
    j=doc.add_paragraph(f"Próximos eventos recomendados para el equipo de {persona.get('Socio', 'N/D')}", style='CustomTitle2')
    j.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(j, "4570F7")

    if not recomendaciones_positivas.empty:
        for idx, fila in recomendaciones_positivas.iterrows():
            evento_info = próximos_eventos[próximos_eventos["Título"] == fila["Evento"]].iloc[0]

            # Título y Score en negrita
            a = doc.add_paragraph(style='List Bullet')
            a.add_run(evento_info['Título']).bold = True
            a.add_run(f" (Score: {fila['Score_num']:.2f}%)").bold = True


            # Tecnologías
            tecnos = evento_info.get("Tecnología", [])
            if isinstance(tecnos, list):
                tecnos_str = ", ".join(tecnos)
            elif isinstance(tecnos, str):
                tecnos_str = tecnos.strip("[]").replace("'", "").replace('"', '')
            else:
                tecnos_str = str(tecnos)
            p = doc.add_paragraph()
            p.add_run("    ○ Tecnologías del evento: ").bold = True
            p.add_run(tecnos_str)

            # Sector
            sectos = evento_info.get("Sector", [])
            if isinstance(sectos, list):
                sectos_str = ", ".join(sectos)
            elif isinstance(sectos, str):
                sectos_str = sectos.strip("[]").replace("'", "").replace('"', '')
            else:
                sectos_str = str(sectos)
            p = doc.add_paragraph()
            p.add_run("    ○ Sectores del evento: ").bold = True
            p.add_run(sectos_str)

            # Ámbito
            ambis = evento_info.get("Ámbito", [])
            if isinstance(ambis, list):
                ambis_str = ", ".join(ambis)
            elif isinstance(ambis, str):
                ambis_str = ambis.strip("[]").replace("'", "").replace('"', '')
            else:
                ambis_str = str(ambis)
            p = doc.add_paragraph()
            p.add_run("    ○ Ámbitos del evento: ").bold = True
            p.add_run(ambis_str)

            # Ubicación
            p = doc.add_paragraph()
            p.add_run("    ○ Ubicación: ").bold = True
            p.add_run(str(evento_info.get('Ubicación', 'N/D')))

            # Provincia (solo si la ubicación no es 'online')
            ubicacion = str(evento_info.get('Ubicación', '')).strip().lower()
            if ubicacion != 'online':
                p = doc.add_paragraph()
                p.add_run("    ○ Provincia donde se celebrará: ").bold = True
                p.add_run(str(evento_info.get('Provincia', 'N/D')))

            
            from datetime import datetime
            
            import sys

            

            fecha_evento = evento_info.get('Fecha', 'N/D')

            from datetime import datetime
            import sys
            import pandas as pd

            

            fecha_str = formatear_fecha_es(fecha_evento)

            p = doc.add_paragraph()
            p.add_run("    ○ Fecha: ").bold = True
            p.add_run(fecha_str)


            doc.add_paragraph("")  # Línea en blanco entre eventos
    else:
        doc.add_paragraph("No se encontraron eventos relevantes para recomendar en este momento.")
        doc.add_paragraph("")

    from datetime import datetime, timedelta

    # Añadir sección de eventos asistidos en el último año
    p=doc.add_paragraph(f"Eventos a los que ha asistido el equipo de {persona.get('Socio', 'N/D')}", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(p, "4570F7")

    
    if not df_eventos_por_evento.empty: 
        for _, evento in df_evento_socios_por_evento.iterrows():
            p = doc.add_paragraph(style='ListBullet')
    
            # Título del evento en negrita
            run_titulo = p.add_run(str(evento["Título"]))
            run_titulo.bold = True
        
            # Fecha en negrita entre paréntesis (si existe)
            if pd.notna(evento.get("Fecha")):
                fecha_str = f" ({evento['Fecha']})"
                run_fecha = p.add_run(fecha_str)
                run_fecha.bold = True
        
            # Separador y asistentes en texto normal
            asistentes = evento.get("Asistentes", "")
            if asistentes:
                run_asistentes = p.add_run(f" — {asistentes}")
    else:
        doc.add_paragraph(
            f"Las personas de {persona.get('Nombre', 'N/D')} no han asistido a ningún evento organizado por SECPhO todavía.",
            style='Normal'
        )
        

    doc.add_paragraph("")

    titulo = doc.add_paragraph("5. Retos Tecnológicos", style='CustomTitle')
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"Aquí se incluyen los retos tecnológicos más afines a las capacidades de {persona.get('Socio', 'N/D')}, junto con los retos en los que la entidad ya ha mostrado interés o participación")
    # Añadir título centrado con estilo personalizado
    

    # Recomendación de retos
    p=doc.add_paragraph("Recomendación de Retos Tecnológicos Activos", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_paragraph_background(p, "F25830")

    recomendaciones_retos = recomendar_retos_por_perfil_socio(persona, retos_futuros, top_n=10)
    recomendaciones_retos["Score_num"] = recomendaciones_retos["ScoreFinal"] * 100  # numérico
    recomendaciones_retos["Similitud %"] = (recomendaciones_retos["ScoreFinal"] * 100).round(2).astype(str) + "%"  # visual

    # Filtrar eventos con puntuación positiva
    recomendaciones_positivas_retos = recomendaciones_retos[recomendaciones_retos["Score_num"] > 0].head(3)

    if not recomendaciones_positivas_retos.empty:
        for idx, fila in recomendaciones_positivas_retos.iterrows():
            # Obtener entidad emisora del reto
            reto_info = retos_futuros[retos_futuros["Título"] == fila["Título"]].iloc[0]
            empresa_persona = str(persona.get('Socio', 'N/D')).strip()

            # --- Comprobar si la empresa de la persona coincide con entidad emisora ---
            entidad_emisora = reto_info.get("Entidad emisora", "")
            entidades_emisoras = []

            if pd.notna(entidad_emisora):
                if isinstance(entidad_emisora, list):
                    entidades_emisoras = [str(e).strip() for e in entidad_emisora]
                elif isinstance(entidad_emisora, str):
                    entidades_emisoras = [e.strip() for e in entidad_emisora.split(",")]
                else:
                    entidades_emisoras = [str(entidad_emisora).strip()]

            if empresa_persona in entidades_emisoras:
                continue  # Saltar si la empresa de la persona es la emisora del reto

            # --- Comprobar si la empresa de la persona ya ha aplicado ---
            entidades_aplican = reto_info.get("Entidades que aplican", [])
            entidades_aplicantes = []

            if pd.notna(entidades_aplican):
                if isinstance(entidades_aplican, list):
                    entidades_aplicantes = [str(e).strip() for e in entidades_aplican]
                elif isinstance(entidades_aplican, str):
                    entidades_aplicantes = [e.strip() for e in entidades_aplican.split(",")]
                else:
                    entidades_aplicantes = [str(entidades_aplican).strip()]

            if empresa_persona in entidades_aplicantes:
                continue  # Saltar si la empresa de la persona ya ha aplicado

            # --- Si pasó ambas comprobaciones, se muestra el reto ---

            # Título y Score en negrita
            a = doc.add_paragraph()
            a.add_run(reto_info['Num. reto']).bold = True
            a.add_run(f" {reto_info['Título']}").bold = True
            a.add_run(f" (Score de similitud: {fila['Score_num']:.2f}%)").bold = True

            # Descripción
            descripcion = reto_info.get("Descripción", "")
            if not descripcion or pd.isna(descripcion) or str(descripcion).strip() == "":
                descripcion = "Descripción no disponible."

            p = doc.add_paragraph()
            p.add_run("    ○ Descripción: ").bold = True
            p.add_run(descripcion)

            # Sector
            sector = reto_info.get("Sector/es", [])
            if isinstance(sector, list):
                sector_str = ", ".join(sector)
            elif isinstance(sector, str):
                sector_str = sector.strip("[]").replace("'", "").replace('"', '')
            else:
                sector_str = str(sector)

            p = doc.add_paragraph()
            p.add_run("    ○ Sector al que pertenece el reto: ").bold = True
            p.add_run(sector_str)

            # Entidad emisora
            p = doc.add_paragraph()
            p.add_run("    ○ Entidad emisora del reto: ").bold = True
            p.add_run(str(reto_info.get('Entidad emisora', 'N/D')))

            # Entidades que aplican
            if not entidades_aplican:
                entidades_texto = "Ninguna"
            else:
                entidades_texto = str(entidades_aplican)

            p = doc.add_paragraph()
            p.add_run("    ○ Entidades que ya han aplicado: ").bold = True
            p.add_run(entidades_texto)

            # Fecha de cierre
            fecha_cierre = reto_info.get('Fecha cierre', 'N/D')
            fecha_cierre_str = formatear_fecha_es(fecha_cierre)

            p = doc.add_paragraph()
            p.add_run("    ○ Fecha de cierre: ").bold = True
            p.add_run(fecha_cierre_str)

            doc.add_paragraph("")  # Espacio al final
    else:
        doc.add_paragraph("No se han encontrado retos tecnológicos para recomendar a esta persona.")
        doc.add_paragraph("")
            
    # Añadir sección de retos en el último año
    empresa_persona = str(persona.get('Socio', 'N/D')).strip()
    p=doc.add_paragraph(f"Retos tecnológicos emitidos por {empresa_persona} gestionados por SECPHO", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(p, "F25830")

    # Lista para guardar retos de la empresa
    retos_de_su_empresa = []

    for _, fila in retos_pasados.iterrows():
        entidad_emisora = fila["Entidad emisora"]

        entidades = []
        if pd.notna(entidad_emisora):
            if isinstance(entidad_emisora, list):
                entidades = [str(e).strip() for e in entidad_emisora]
            elif isinstance(entidad_emisora, str):
                entidades = [e.strip() for e in entidad_emisora.split(",")]
            else:
                entidades = [str(entidad_emisora).strip()]
        
        if empresa_persona in entidades:
            retos_de_su_empresa.append(fila)

    # Mostrar resultados
    if not retos_de_su_empresa:
        doc.add_paragraph(f"Por ahora,{empresa_persona} no ha emitido retos tecnológicos a través de SECPHO. Nuestro equipo puede ayudarte tanto en la definición y dinamización de retos como en la búsqueda de partners mediante tech scouting que aporten soluciones concretas.")
        doc.add_paragraph("")
    else:
        for fila in retos_de_su_empresa:
            p = doc.add_paragraph()
            p.add_run(f"{fila['Num. reto']} ").bold = True
            p.add_run(f"{fila['Título']}").bold = True


            # Descripción
            descripcion = fila.get("Descripción", "")
            if not descripcion or pd.isna(descripcion) or str(descripcion).strip() == "":
                descripcion = "Descripción no disponible."
            p = doc.add_paragraph()
            p.add_run("    ○ Descripción: ").bold = True
            p.add_run(descripcion)


            # Sector
            sector = fila["Sector/es"]
            if isinstance(sector, list):
                sector_str = ", ".join(sector)
            elif isinstance(sector, str):
                sector_str = sector.strip("[]").replace("'", "").replace('"', '')
            else:
                sector_str = str(sector)

            p = doc.add_paragraph()
            p.add_run("    ○ Sector al que pertenece el reto: ").bold = True
            p.add_run(sector_str)

            # Entidades que aplican
            entidades_aplican = fila.get('Entidades que aplican', None)
            entidades_texto = "Ninguna" if not entidades_aplican else str(entidades_aplican)

            p = doc.add_paragraph()
            p.add_run("    ○ Entidades que ya han aplicado: ").bold = True
            p.add_run(entidades_texto)

            # Fecha de cierre
            fecha_cierre = fila['Fecha cierre']
            fecha_cierre_str = formatear_fecha_es(fecha_cierre)
            p = doc.add_paragraph()
            p.add_run("    ○ Fecha de cierre: ").bold = True
            p.add_run(fecha_cierre_str)
            doc.add_paragraph("")


    # Añadir sección de retos en el último año
    empresa_persona = str(persona.get('Socio', 'N/D')).strip()
    p=doc.add_paragraph(f"Retos tecnológicos a los que {empresa_persona} ha aplicado", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    set_paragraph_background(p, "F25830")

    # Lista para guardar retos de la empresa
    retos_de_su_empresa = []

    for _, fila in retos_pasados.iterrows():
        entidades_aplicantes = fila["Entidades que aplican"]

        entidades = []
        if pd.notna(entidades_aplicantes):
            if isinstance(entidades_aplicantes, list):
                entidades = [str(e).strip() for e in entidades_aplicantes]
            elif isinstance(entidades_aplicantes, str):
                entidades = [e.strip() for e in entidades_aplicantes.split(",")]
            else:
                entidades = [str(entidades_aplicantes).strip()]
        
        if empresa_persona in entidades:
            retos_de_su_empresa.append(fila)

    # Mostrar resultados
    if not retos_de_su_empresa:
        doc.add_paragraph(f"{empresa_persona} nunca ha aplicado a un reto tecnológico.")
        doc.add_paragraph("")
    else:
        for fila in retos_de_su_empresa:
            p = doc.add_paragraph()
            p.add_run(f"{fila['Num. reto']} ").bold = True
            p.add_run(f"{fila['Título']}").bold = True


            # Descripción
            descripcion = fila.get("Descripción", "")
            if not descripcion or pd.isna(descripcion) or str(descripcion).strip() == "":
                descripcion = "Descripción no disponible."
            p = doc.add_paragraph()
            p.add_run("    ○ Descripción: ").bold = True
            p.add_run(descripcion)


            # Sector
            sector = fila["Sector/es"]
            if isinstance(sector, list):
                sector_str = ", ".join(sector)
            elif isinstance(sector, str):
                sector_str = sector.strip("[]").replace("'", "").replace('"', '')
            else:
                sector_str = str(sector)

            p = doc.add_paragraph()
            p.add_run("    ○ Sector al que pertenece el reto: ").bold = True
            p.add_run(sector_str)

            # Entidad emisora
            emisor = fila["Entidad emisora"]
            if isinstance(emisor, list):
                emisor_str = ", ".join(emisor)
            elif isinstance(emisor, str):
                emisor_str = emisor.strip("[]").replace("'", "").replace('"', '')
            else:
                emisor_str = str(emisor)

            p = doc.add_paragraph()
            p.add_run("    ○ Entidad emisora del reto: ").bold = True
            p.add_run(emisor_str)

            # Entidades que aplican
            entidades_aplican = fila.get('Entidades que aplican', None)
            entidades_texto = "Ninguna" if not entidades_aplican else str(entidades_aplican)

            p = doc.add_paragraph()
            p.add_run("    ○ Entidades que ya han aplicado: ").bold = True
            p.add_run(entidades_texto)

            # Fecha de cierre
            fecha_cierre = fila['Fecha cierre']
            fecha_cierre_str = formatear_fecha_es(fecha_cierre)
            p = doc.add_paragraph()
            p.add_run("    ○ Fecha de cierre: ").bold = True
            p.add_run(fecha_cierre_str) 
            doc.add_paragraph("")   


    # Añadir título centrado con estilo personalizado
    titulo = doc.add_paragraph("6. Proyectos", style='CustomTitle')
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph(f"Este apartado recoge tanto los proyectos en los que {persona.get('Socio', 'N/D')} ha colaborado como aquellos coordinados por secpho que resultan de interés estratégico por su alineación con las capacidades, tecnologías y sectores de interés de {persona.get('Socio', 'N/D')}")
    
    # Añadir sección de proyectos de la empresa
    empresa_persona = str(persona.get('Socio', 'N/D')).strip()
    p=doc.add_paragraph(f"Proyectos en los que {empresa_persona} ha  colaborado como Partner", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_paragraph_background(p, "4C49F2")

    proyectos_filtrados = proyectos_de_socio(proyectos, empresa_persona) 
    if proyectos_filtrados.empty:
        doc.add_paragraph(f"Actualmente, {empresa_persona} no ha participado como partner en proyectos a través de secpho. Estamos a tu disposición para identificar oportunidades y conectar con entidades afines que puedan convertirse en socios estratégicos en futuras iniciativas. ")
        doc.add_paragraph("")
    else:
        for _, fila in proyectos_filtrados.iterrows():
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{fila['Título de la Idea/Proyecto']}").bold = True



            # Fecha de inicio
            fecha_inicio = fila[' Inicio']
            fecha_inicio_str = formatear_fecha_es(fecha_inicio)
            p = doc.add_paragraph()
            p.add_run("    ○ Fecha de inicio del proyecto: ").bold = True
            p.add_run(fecha_inicio_str)

            # Fecha de cierre
            fecha_final = fila[' Final']
            fecha_final_str = formatear_fecha_es(fecha_final)
            p = doc.add_paragraph()
            p.add_run("    ○ Fecha de finalización del proyecto: ").bold = True
            p.add_run(fecha_final_str)

            # partners
            part = fila["Partners"]
            if isinstance(part, list):
                part_str = ", ".join(part)
            elif isinstance(part, str):
                part_str = part.strip("[]").replace("'", "").replace('"', '')
            else:
                part_str = str(part)
            if part_str == "":
                part_str = 'No definidos'
            p = doc.add_paragraph()
            p.add_run("    ○ Partners del proyecto: ").bold = True
            p.add_run(part_str)

            # Origen fondos
            p = doc.add_paragraph()
            p.add_run("    ○ Origen de los fondos: ").bold = True
            p.add_run(str(fila.get('Origen fondos ', 'N/D')))

            # Programa financiación
            programa = fila.get("Programa financiacion ", "")
            if not programa or pd.isna(programa) or str(programa).strip() == "":
                programa = "No consta"
            p = doc.add_paragraph()
            p.add_run("    ○ Programa de financiación: ").bold = True
            p.add_run(programa)

            # Presupuesto total
            presu = fila.get("Presupuesto total (€)", "")
            if pd.notna(presu):
                p = doc.add_paragraph()
                p.add_run("- Presupuest total: ").bold = True
                p.add_run(f"{presu:,.2f} €")  
            else:
                p = doc.add_paragraph()
                p.add_run("    ○ Presupuest total: ").bold = True
                p.add_run("No consta")


            


            # tecnologias
            tecno = fila["Tecnologías"]
            if isinstance(tecno, list):
                tecno_str = ", ".join(tecno)
            elif isinstance(tecno, str):
                tecno_str = tecno.strip("[]").replace("'", "").replace('"', '')
            else:
                tecno_str = str(tecno)
            if tecno_str == "":
                tecno_str = 'No definidas'
            p = doc.add_paragraph()
            p.add_run("    ○ Tecnologías del proyecto: ").bold = True
            p.add_run(tecno_str)

            # Sector
            sector = fila["Sectores"]
            if isinstance(sector, list):
                sector_str = ", ".join(sector)
            elif isinstance(sector, str):
                sector_str = sector.strip("[]").replace("'", "").replace('"', '')
            else:
                sector_str = str(sector)
            if sector_str == "":
                sector_str = 'No definido'
            p = doc.add_paragraph()
            p.add_run("    ○ Sectores del proyecto: ").bold = True
            p.add_run(sector_str)

            # ambitos
            ambi = fila["Ámbitos"]
            if isinstance(ambi, list):
                ambi_str = ", ".join(ambi)
            elif isinstance(ambi, str):
                ambi_str = ambi.strip("[]").replace("'", "").replace('"', '')
            else:
                ambi_str = str(ambi)
            if ambi_str == "":
                ambi_str = 'No definido'
            p = doc.add_paragraph()
            p.add_run("    ○ Ámbitos del proyecto: ").bold = True
            p.add_run(ambi_str)
            doc.add_paragraph("")

    tecnologias_persona = persona.get('Tecnologías')
    sectores_persona = persona.get('Sectores')

    proyectos_rel = proyectos_relacionados(proyectos, tecnologias_persona, sectores_persona)

    p=doc.add_paragraph(f"Histórico de proyectos coordinados por SECPHO en ámbitos de interés para {persona.get('Socio', 'N/D')}", style='CustomTitle2')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_paragraph_background(p, "4C49F2") 

    if proyectos_rel.empty:
        doc.add_paragraph(f"No se ha realizado ningún proyecto de interés sobre el que informar a miembros de {persona.get('Socio', 'N/D')}")
        doc.add_paragraph("")
    else:
        def esta_vacio(x):
            if x is None:
                return True
            if isinstance(x, float) and np.isnan(x):
                return True
            if isinstance(x, str):
                s = x.strip()
                return (s == "" or s.lower() in {"nan", "n/d", "nd", "none"})
            if isinstance(x, (list, tuple, set)):
                return len(x) == 0
            return False

        def fmt_valor(v):
            if isinstance(v, (list, tuple, set)):
                return ", ".join(str(x) for x in v)
            return str(v)

        for _, fila in proyectos_rel.iterrows():
            tec = fila.get('Tecnología de coincidencia')
            sec = fila.get('Sector de coincidencia')

            partes = []
            if not esta_vacio(tec):
                partes.append(f"Tecnología: {fmt_valor(tec)}")
            if not esta_vacio(sec):
                partes.append(f"Sector: {fmt_valor(sec)}")

            # Si no hay nada que mostrar, saltar
            if not partes:
                continue

            p = doc.add_paragraph(" || ".join(partes), style='CustomTitle2')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            set_paragraph_background(p, "8C8AFF")

    
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{fila['Título de la Idea/Proyecto']}").bold = True



            # Fecha de inicio
            fecha_inicio = fila[' Inicio']
            fecha_inicio_str = formatear_fecha_es(fecha_inicio)
            p = doc.add_paragraph()
            p.add_run("    ○ Fecha de inicio del proyecto: ").bold = True
            p.add_run(fecha_inicio_str)

            # Fecha de cierre
            fecha_final = fila[' Final']
            fecha_final_str = formatear_fecha_es(fecha_final)
            p = doc.add_paragraph()
            p.add_run("    ○ Fecha de finalización del proyecto: ").bold = True
            p.add_run(fecha_final_str)

            # partners
            part = fila["Partners"]
            if isinstance(part, list):
                part_str = ", ".join(part)
            elif isinstance(part, str):
                part_str = part.strip("[]").replace("'", "").replace('"', '')
            else:
                part_str = str(part)
            if part_str == "":
                part_str = 'No definidos'
            p = doc.add_paragraph()
            p.add_run("    ○ Partners del proyecto: ").bold = True
            p.add_run(part_str)

            # Origen fondos
            p = doc.add_paragraph()
            p.add_run("    ○ Origen de los fondos: ").bold = True
            p.add_run(str(fila.get('Origen fondos ', 'N/D')))

            # Programa financiación
            programa = fila.get("Programa financiacion ", "")
            if not programa or pd.isna(programa) or str(programa).strip() == "":
                programa = "No consta"
            p = doc.add_paragraph()
            p.add_run("    ○ Programa de financiación: ").bold = True
            p.add_run(programa)

            # Presupuesto total
            presu = fila.get("Presupuesto total (€)", "")
            if pd.notna(presu):
                p = doc.add_paragraph()
                p.add_run("    ○ Presupuest total: ").bold = True
                p.add_run(f"{presu:,.2f} €")  
            else:
                p = doc.add_paragraph()
                p.add_run("- Presupuest total: ").bold = True
                p.add_run("No consta")


            


            # tecnologias
            tecno = fila["Tecnologías"]
            if isinstance(tecno, list):
                tecno_str = ", ".join(tecno)
            elif isinstance(tecno, str):
                tecno_str = tecno.strip("[]").replace("'", "").replace('"', '')
            else:
                tecno_str = str(tecno)
            if tecno_str == "":
                tecno_str = 'No definidas'
            p = doc.add_paragraph()
            p.add_run("    ○ Tecnologías del proyecto: ").bold = True
            p.add_run(tecno_str)

            # Sector
            sector = fila["Sectores"]
            if isinstance(sector, list):
                sector_str = ", ".join(sector)
            elif isinstance(sector, str):
                sector_str = sector.strip("[]").replace("'", "").replace('"', '')
            else:
                sector_str = str(sector)
            if sector_str == "":
                sector_str = 'No definido'
            p = doc.add_paragraph()
            p.add_run("    ○ Sectores del proyecto: ").bold = True
            p.add_run(sector_str)

            # ambitos
            ambi = fila["Ámbitos"]
            if isinstance(ambi, list):
                ambi_str = ", ".join(ambi)
            elif isinstance(ambi, str):
                ambi_str = ambi.strip("[]").replace("'", "").replace('"', '')
            else:
                ambi_str = str(ambi)
            if ambi_str == "":
                ambi_str = 'No definido'
            p = doc.add_paragraph()
            p.add_run("    ○ Ámbitos del proyecto: ").bold = True
            p.add_run(ambi_str)
            doc.add_paragraph("")




    # Guardar en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
#------------------------------------------------------------------------------------------------------------


# 🎨 Interfaz de Streamlit
import base64
from io import BytesIO
import streamlit as st
from PIL import Image

st.set_page_config(
        page_title="Generador de Informes",
        page_icon="📝",
        layout="wide"  # 👈 Ocupa toda la pantalla
    )

def check_password():
    def password_entered():
        if st.session_state["password"] == "MARTA":
            st.session_state["password_correct"] = True
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        # Primer acceso
        st.text_input("Contraseña", type="password", on_change=password_entered, key="password")
        return False
    elif not st.session_state["password_correct"]:
        # Contraseña incorrecta
        st.text_input("Contraseña", type="password", on_change=password_entered, key="password")
        st.error("Contraseña incorrecta ❌")
        return False
    else:
        # Contraseña correcta ✅
        return True

if check_password():
    # Mostrar imagen con st.image (no en HTML)
    st.image("logo.png", width=150)

    # 💅 Estilos personalizados para modo oscuro
    st.markdown("""
        <style>
            /* Oculta menú y pie de página */
            #MainMenu, footer {visibility: hidden;}

            /* Aumenta el ancho y márgenes */
            .stApp {
                max-width: 100%;
                padding: 2rem 5rem;
                background-color: #0e1117;
                color: white;
            }

            /* Aumenta contraste de encabezados */
            h1, h2, h3 {
                color: white;
            }
        </style>
    """, unsafe_allow_html=True)

    # Contenido principal
    st.markdown(
        "<h2 style='text-align:center;'>📝 Generador de Informes de Valor y Oportunidades</h2>",
        unsafe_allow_html=True
    )


    st.markdown(
        "<h2 style='font-size:30px; color:#4c49f2; text-align:center;'>¿Qué informe quieres generar?</h2>",
        unsafe_allow_html=True
    )

    from unidecode import unidecode

    # --- Estado persistente ---
    if "modo" not in st.session_state:
        st.session_state.modo = None   # None | "persona" | "empresa"

    def set_modo(m):
        st.session_state.modo = m

    def normalizar_nombre(nombre: str) -> str:
        return unidecode(str(nombre)).lower().replace(" ", "")

    # --- Botones centrados ---
    col1, col2, col3, col4, col5, col6 = st.columns([2,1,1,1,1,2])
    with col3:
        st.button("👤 Persona", key="btn_persona", on_click=set_modo, args=("persona",))
    with col4:
        st.button("🏢 Socio", key="btn_empresa", on_click=set_modo, args=("empresa",))


    # --- SOLO RENDERIZAR FORMULARIOS SI HAY MODO ---
    if st.session_state.modo == "persona":
        st.markdown("Selecciona una persona para generar un informe personalizado ✨")

        nombres_completos = [f"{row['Nombre']} {row['Apellidos']}" for _, row in miembros.iterrows()]

        informe = None
        nombre_real = None

        with st.form("form_persona"):
            nombre_seleccionado = st.selectbox("👤 Selecciona una persona", sorted(nombres_completos))
            submitted = st.form_submit_button("🎯 Generar informe")

            if submitted and nombre_seleccionado:
                nombre_normalizado = normalizar_nombre(nombre_seleccionado)
                fila_match = miembros[
                    miembros.apply(
                        lambda row: normalizar_nombre(f"{row['Nombre']} {row['Apellidos']}") == nombre_normalizado,
                        axis=1
                    )
                ]
                if fila_match.empty:
                    st.error("❌ No se encontró la persona seleccionada en el DataFrame.")
                else:
                    nombre_real = fila_match.iloc[0]["Nombre completo"]
                    informe = generar_informe_persona(nombre_real)
                    st.success(f"✅ Informe generado para **{nombre_real}**")

        # 👉 Fuera del formulario
        if informe and nombre_real:
            st.download_button(
                label="📥 Descargar informe Word",
                data=informe,
                file_name=f"Informe_{nombre_real.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


    elif st.session_state.modo == "empresa":
        st.markdown("Selecciona un socio para generar un informe de empresa ✨")

        empresas = socios["Socio"].dropna().unique().tolist()

        # Formulario solo para seleccionar y enviar
        with st.form("form_empresa"):
            empresa_seleccionada = st.selectbox("🏢 Selecciona un socio", sorted(empresas))
            submitted_e = st.form_submit_button("🎯 Generar informe socio")

        # ✅ Fuera del formulario
        if submitted_e and empresa_seleccionada:
            nombre_normalizado = normalizar_nombre(empresa_seleccionada)
            fila_match = socios[
                socios.apply(lambda row: normalizar_nombre(row.get("Socio", "")) == nombre_normalizado, axis=1)
            ]
            if fila_match.empty:
                st.error("❌ No se encontró el socio seleccionado en el DataFrame.")
            else:
                nombre_real = fila_match.iloc[0]["Nombre completo"]
                informe = generar_informe_socio(nombre_real)
                st.success(f"✅ Informe generado para **{nombre_real}**")

                # Download button fuera del formulario
                st.download_button(
                    label="📥 Descargar informe Word",
                    data=informe,
                    file_name=f"Informe_{nombre_real.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )


